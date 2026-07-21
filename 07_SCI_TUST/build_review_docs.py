#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""build_review_docs.py -- 產生 07_SCI_TUST 的 Word 審閱文件樹（Wade 於 Word 中批改）。

結構（Wade 07-18 指令）：
  Chapter00_總綱/   題目賣點材料 / 圖表總覽 / 參考文獻總集APA / 摘要關鍵字 / 碩論濃縮對照
  Chapter01~06/     子節架構/ 段落核心/ Figure/ Table/ Reference/  各一 docx
冪等：重跑即重建全部 docx（Wade 的批改請另存 _批 版本，或直接改在原檔上告知）。
格式：中文標楷體＋英數 Times New Roman；⚖ 待裁決段落以【⚖ 待裁決】前綴＋粗體標示。
"""
import io
import sys
from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
ROOT = Path(__file__).parent


def new_doc():
    d = docx.Document()
    st = d.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(12)
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
    return d


def set_font(run, size=12, bold=False, color=None):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rpr = run._r.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        from docx.oxml import OxmlElement
        rf = OxmlElement("w:rFonts")
        rpr.insert(0, rf)
    rf.set(qn("w:ascii"), "Times New Roman")
    rf.set(qn("w:hAnsi"), "Times New Roman")
    rf.set(qn("w:eastAsia"), "標楷體")


def H(d, text, lv=1):
    p = d.add_paragraph()
    r = p.add_run(text)
    set_font(r, size=16 if lv == 1 else 14, bold=True)
    p.space_after = Pt(6)
    return p


def P(d, text, size=12, bold=False, indent=False):
    p = d.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Pt(18)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold)
    return p


def DEC(d, text):
    """⚖ 待裁決 block -- red bold prefix, roomy for comments."""
    p = d.add_paragraph()
    r1 = p.add_run("【⚖ 待裁決】")
    set_font(r1, bold=True, color=(0xC0, 0x00, 0x00))
    r2 = p.add_run(text)
    set_font(r2)
    p.paragraph_format.space_after = Pt(12)
    return p


def TBL(d, header, rows, widths=None):
    t = d.add_table(rows=1 + len(rows), cols=len(header))
    t.style = "Table Grid"
    for j, h in enumerate(header):
        c = t.rows[0].cells[j]
        c.text = ""
        r = c.paragraphs[0].add_run(h)
        set_font(r, size=11, bold=True)
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            c = t.rows[1 + i].cells[j]
            c.text = ""
            r = c.paragraphs[0].add_run(str(v))
            set_font(r, size=10.5)
    return t


def save(d, *parts):
    p = ROOT.joinpath(*parts)
    p.parent.mkdir(parents=True, exist_ok=True)
    d.save(p)
    print("OK", p.relative_to(ROOT))


# ============================================================================
# Chapter00 總綱
# ============================================================================
d = new_doc()
H(d, "文章題目、賣點、材料（框架首頁 v1，2026-07-18）")
P(d, "說明：本頁為全文之錨，仿 IJRMMS 前作的框架首頁迭代法——此頁穩定前不動正文。"
     "請直接在 Word 中批改、圈選、加註。", size=11)
H(d, "一、題目（v2，Wade 07-19 裁示後修正：案例＋機制入題，方法字全退回賣點層）", 2)
P(d, "構式範本＝團隊自家 Q1：Chiu, Lee & Wang (2017, TUST, 被引94)「Lining crack "
     "evolution of an operational tunnel influenced by slope instability」＝"
     "【現象+案例】influenced by【原因】，零方法字。", size=11)
P(d, "N-A（系列構式直接繼承，換驅動源——首選）：", bold=True)
P(d, "Lining crack evolution of an operating mountain railway tunnel influenced by "
     "cyclic groundwater fluctuation", indent=True)
P(d, "營運山岳鐵路隧道襯砌裂縫受地下水位循環升降影響之演化", indent=True)
P(d, "與 2017 前作形成明確系列感（slope instability → cyclic groundwater "
     "fluctuation）。", size=10.5, indent=True)
P(d, "N-B（機制入題，接前作「滯動」術語＝理論＋案例雙系列繼承）：", bold=True)
P(d, "Lining cracking of an operating mountain railway tunnel driven by groundwater-"
     "fluctuation-induced intermittent time-dependent deformation", indent=True)
P(d, "地下水位升降引致滯動（起閉）依時變形驅動之營運山岳鐵路隧道襯砌開裂", indent=True)
P(d, "N-C（現象前置＋機制副標）：", bold=True)
P(d, "Lining crack evolution in an operating mountain railway tunnel under cyclic "
     "groundwater fluctuation: intermittent time-dependent deformation as the "
     "driving mechanism", indent=True)
DEC(d, "N-A / N-B / N-C？（我推 N-A：最短、系列感最強；56 篇 Q1 題目解剖統計出爐後"
       "會附驗證數據）")
H(d, "二、賣點（5 項，按主打順序）", 2)
for i, s in enumerate([
    "跨尺度因果鏈完整量化：地下水年循環（坡地尺度）→近場滲流→門檻滯動潛變→襯砌離散開裂"
    "（單一裂縫級）；三尺度巢狀地質模型錨定＋一致性資料傳遞。文獻中隧道–邊坡互制幾乎"
    "全為地震題（Xin 2024; Tian 2026, 皆 TUST），水文驅動跨尺度鏈為空白。",
    "滯動機制之襯砌尺度延伸：前作（IJRMMS, in review）建立圍岩尺度滯動理論，本文推進到"
    "「滯動→襯砌損傷節律」：雨峰爆發 A_wet=7.0、退水凍結 A_frz=0.0046、濕窗貢獻 91% 損傷。",
    "BPM 離散襯砌之裂縫型態學：2.08M 鍵結無配筋 RC 襯砌，重現裂縫帶狀分區＋三軌型態分類，"
    "並與現地 LiDAR 展開圖同座標對照。",
    "雙向損傷回饋初探（放討論章）：D(s,y)→E 割線勁度回饋之 5 天交替分區耦合，day-130 損傷"
    "達單向 2.83×；力回饋過定性論證（Picard 發散實證）本身即方法貢獻。",
    "維養決策視角（TUST 甜蜜點）：真實營運隧道＋分級調查制度＋水位監測 → "
    "「地下水位管理＝襯砌損傷管理」。"], 1):
    P(d, f"賣點 {i}：{s}", indent=True)
DEC(d, "五條賣點之措辭與排序？（前作經驗：賣點是被改最多次的元素，請大方下刀）")
H(d, "三、材料（5 項食材）", 2)
for i, s in enumerate([
    "修正力學模式（門檻黏彈塑 Burgers-Mohr，λ≡T 開關；s1 門檻掃描定準法）——引前作、只寫差異",
    "三維工程地質模型（坡地 4 層／隧道 6 層／互制模型；340 m 鑽探＋岩心＋室內試驗給出處"
    "——202411 附錄 11/12/13）",
    "真實案例資料（#46 隧道 1233 m、R=50 曲線：LiDAR 全斷面兩期、三期異狀展開圖、"
    "裂縫計 110–112 持續緩剪、112 年修復履歷、0403 地震後新裂縫）",
    "單向弱耦合定版成果（05 v6：11 階段 130 天、37,980 斷鍵、環199/斜19/縱3 m、"
    "外壓 946 kPa、左肩推力 1,029 kN/m）",
    "雙向強耦合鏈（06 T5：26×5 天交替、D→E 回饋、107,479 斷鍵=2.83×、逐 tick 全紀錄）"], 1):
    P(d, f"材料 {i}：{s}", indent=True)
H(d, "四、討論要點（5 項）", 2)
for s in ["雙向 vs 單向放大之組成（回饋效應 vs 細時距路徑效應；L0 對照鏈拆解）",
          "門檻 T 的物理意義與定準敏感度（回應口委黃老師；與前作 λ≈0.4 之尺度差異論述）",
          "水位擺幅情境（實測 ~30 m vs 模擬 100 m 之情境放大論證）",
          "f=0.25 縮尺與趨勢級解讀邊界",
          "水位管理＝損傷管理（滯動泵概念收尾）"]:
    P(d, "・" + s, indent=True)
save(d, "Chapter00_總綱", "00_題目賣點材料.docx")

d = new_doc()
H(d, "全文圖表總覽（15 圖＋5 表；上限 16 圖）")
TBL(d, ["圖", "內容與故事", "素材／加工", "狀態"], [
    ["Fig 1", "場址與案例綜覽：林鐵×曲線隧道×坡地系統空間關係", "附圖95 圈繪＋附圖96 剖面＋App5 幾何", "新繪"],
    ["Fig 2", "#46 現地病徵時序：三期展開圖＋裂縫照片＋裂縫計緩剪", "附圖93/97/100–102＋裂縫計", "重組英化"],
    ["Fig 3", "水文驅動：水位計年擺幅→W 面家族→模擬情境", "碩論表4-1＋擬合圖（⚖需主報告時序）", "改繪"],
    ["Fig 4", "跨尺度方法流程＋資料傳遞（虛線預告 D→E）", "圖5-01 擴充", "改造"],
    ["Fig 5", "三尺度模型三面板", "圖5-02+03+04 合併", "合併"],
    ["Fig 6", "τ-p 門檻機制：濕季應力雲移入門檻帶", "圖5-14", "現成"],
    ["Fig 7", "門檻活化演化：乾衰減→雨爆發→退凍結", "圖5-08+5-12 縮幀合併", "合併"],
    ["Fig 8", "近場排水暈", "圖5-11", "現成"],
    ["Fig 9", "襯砌損傷史 A_wet/A_frz", "圖5-15", "現成"],
    ["Fig 10", "襯砌外壓＋內力：946 kPa 帶＋1,029 kN/m 力偶", "圖5-16+5-17 合併", "合併"],
    ["Fig 11", "裂縫發展與分類（與 Fig2 同座標展開設計）", "圖5-19 改造", "改造"],
    ["Fig 12", "三維裂縫演化（hero）", "圖5-20", "現成"],
    ["Fig 13", "雙向耦合方法＋單雙向時間線", "FIG06-01 期刊化＋小示意", "改造"],
    ["Fig 14", "D(s,y) 損傷圖演化", "FIG06-02 期刊化", "改造"],
    ["Fig 15", "滯動損傷泵概念圖（bookend）", "新繪（pptx 示意，前作 Fig12 手法）", "新繪"],
])
P(d, "")
TBL(d, ["表", "內容", "備註"], [
    ["Table 1", "文獻案例彙編（Author/Case/Lithology/Model/Method/Focus/Issue/Highlight 8 欄）", "中性、不含 this study 行（前作教訓）"],
    ["Table 2", "三尺度模型組態與參數（含出處欄）", "回應 K0/參數來源問"],
    ["Table 3", "門檻與潛變參數（定準依據＋物理意義欄）", "回應門檻依據問"],
    ["Table 4", "現地調查監測資料彙整（LiDAR 兩期/裂縫計/鑽探/試驗）", "案例可信度"],
    ["Table 5", "單向 vs 雙向逐階段對照（含 QA gates 摘要）", "討論章"],
])
DEC(d, "15 圖可否？（可併 Fig7+8 成 14 張；上限 16）表 5 張可否？")
save(d, "Chapter00_總綱", "01_圖表總覽.docx")

d = new_doc()
H(d, "參考文獻總集（APA 7；v2＝主庫 REFS_MASTER 全量）")
P(d, "主庫 56 篇（TUST 31 篇=55%，目標 ≥50% ✓）；全數 Crossref 逐 DOI 驗證＋即時"
     "被引數、零 MDPI。⚖ 略超 40–50 目標：請圈刪 5–8 篇（候選=L3/L6 低引或低階期刊），"
     "或授權依「引用×相關×期刊等級」自裁。", size=11)
_master = (ROOT / "refs" / "REFS_MASTER.md").read_text(encoding="utf-8").splitlines()
for _ln in _master:
    if _ln.startswith("## "):
        H(d, _ln[3:], 2)
    elif _ln.startswith("- ["):
        P(d, _ln[2:].replace("**[TUST]**", "[TUST]"), size=10.5, indent=True)
save(d, "Chapter00_總綱", "02_參考文獻總集_APA.docx")

d = new_doc()
H(d, "摘要與關鍵字（v0 草案）")
P(d, "格式依前作：單段 ~450 字元，問題→缺口→方法→案例→「結果顯示」＋三個分號結論。"
     "數字待 L0 對照與終版圖表定案後回填（【】為佔位）。", size=11)
P(d, "營運中山岳隧道之襯砌開裂與地下水活動密切相關，然而現行依時變形研究多止於圍岩尺度，"
     "且隧道與坡地系統之互制分析多僅考慮地震荷載，長期水文循環驅動之襯砌損傷機制尚缺乏"
     "跨尺度之量化途徑。本文以阿里山林業鐵路某曲線段隧道為案例，建立坡地尺度滲流—隧道"
     "尺度門檻滯動潛變—襯砌尺度離散開裂之三尺度單向傳遞鏈，並以損傷—勁度回饋之交替耦合"
     "初探襯砌劣化之自增強效應；模型以三維工程地質模型錨定，材料參數具鑽探與室內試驗"
     "出處，模擬結果與現地雷射掃描裂縫展開圖同座標對照。結果顯示：雨峰窗口貢獻約九成之"
     "襯砌損傷增量，濕季損傷速率為初始乾季之 7.0 倍、退水後降至雨季之 0.5%，水位相位"
     "實質控制損傷節律；模擬裂縫呈帶狀分區且以環向為主導，與曲線異狀段之現地展開圖對照"
     "【吻合度敘述待 W1 統計】；雙向損傷回饋使 130 天累積損傷達單向之 2.83 倍，顯示"
     "襯砌劣化評估若忽略互制回饋將顯著低估【比例待 L0 分解】。")
H(d, "關鍵字（5–8）", 2)
P(d, "中：跨尺度數值模擬；襯砌開裂；地下水位循環；滯動潛變；FDM–DEM 耦合；隧道維養")
P(d, "EN: cross-scale simulation; lining cracking; groundwater cycling; "
     "intermittent creep; FDM–DEM coupling; tunnel maintenance")
DEC(d, "摘要走向與關鍵字組合？（「損傷泵」一詞是否入摘要？）")
save(d, "Chapter00_總綱", "03_摘要與關鍵字.docx")

d = new_doc()
H(d, "碩論→期刊濃縮對照表（第一刀：解析解全砍）")
TBL(d, ["碩論內容", "期刊處置", "說明"], [
    ["Ch1 緒論", "濃縮入前言 P1/P5", "研究方法沿革敘述刪；動機直接接工程問題"],
    ["Ch2 文獻回顧—解析解相關", "【砍】", "Wade 指令；僅 Fahimifar 2009 留作討論章量級對照一句"],
    ["Ch2 文獻回顧—潛變理論/模型系譜", "砍，引前作代替", "期刊前言以 40–50 篇文獻重建，不做回顧章"],
    ["Ch3 案例（TT 版）", "濃縮入期刊 Ch3", "保留病徵分類＋監測時序；沿革/制度細節刪"],
    ["Ch4 多尺度地質模型", "濃縮為期刊 3.3 一節＋Table 2", "建模過程細節刪；只留組態＋參數出處"],
    ["Ch5.1–5.2 方法（三模型/傳遞）", "濃縮入期刊 Ch2", "Kabsch/門檻/BPM 各一段；QA 一段"],
    ["Ch5.3–5.4 單向成果", "期刊 Ch4（圖 21→9 張）", "刪過程性圖（5-06/07/09/10/13）；⚖ 5-10 驅動場驗證圖是否保留？"],
    ["Ch5.5 侷限", "拆入討論 5.2＋結論限制", "——"],
    ["Ch6 結論", "重寫（5 貢獻＋2 限制）", "不復述前文"],
    ["（碩論無）雙向耦合", "期刊 5.1 新寫", "06 成果＋初探定位"],
])
DEC(d, "還要濃縮/砍什麼？候選：①坡地尺度成果篇幅（4.1 一段 vs 完整小節）"
       "②大模驅動場驗證圖 5-10 ③QA/gates 寫入深度（一段 vs 表格化入 Table 5）")
save(d, "Chapter00_總綱", "04_碩論濃縮對照表.docx")

# ============================================================================
# Chapter01~06
# ============================================================================
CH = {
 "Chapter01_前言": dict(
  arch=[("段落結構（5 段）", None),
        ("P1 研究背景：營運山岳隧道老化×極端水文＝維養挑戰；台灣山岳鐵路案例特殊性（文資＋營運中）", "L5＋Wang TT 2024"),
        ("P2 近期研究（水文→襯砌＋依時變形）：降雨/地下水對襯砌載重與劣化；潛變本構與依時變形進展", "L1＋L2"),
        ("P3 近期研究（數值方法）：連續-離散耦合模擬襯砌開裂之發展；跨尺度方法", "L3"),
        ("P4 研究瓶頸（然而…）：①依時變形止於圍岩尺度②隧道–邊坡互制只做地震③水文當靜態工況；皆單向無回饋", "L4 收口"),
        ("P5 目的與貢獻（不同於…）：三層貢獻宣告＋路線圖", "—")],
  core=["P1 核心句：襯砌裂縫不是靜態缺陷，而是水文循環驅動下持續演化的損傷過程——營運中山岳隧道的維養須理解其節律。",
        "P2 核心句：現有研究已確認水對襯砌載重與依時變形的個別影響，但多為單尺度、單向、"
        "靜態工況之處理。",
        "P3 核心句：連續-離散耦合已能模擬襯砌開裂，然與坡地尺度水文過程之銜接尚無先例。",
        "P4 缺口句（然而）：三缺口逐一點名，各以文獻收尾。",
        "P5 貢獻句（不同於…突破…）：跨尺度鏈＋滯動襯砌延伸＋雙向回饋初探。"],
  fig="本章無圖（Table 1 文獻彙編置於 P2/P3 之間）。",
  tab="Table 1 文獻案例彙編：8 欄中性矩陣（Author/Case/Lithology/Model/Method/Focus/Issue/Highlight），約 10–12 列；不含 this study 行。",
  ref="P1←L5(Moradi/Ou/Chen J)＋Wang TT 2024；P2←L1(An/Chang)＋L2(Zhang/Mao/Wang X＋前作)；P3←L3(Yan/Wang Z/Zheng)；P4←L4(Xin/Tian Y/Tian X)。"),
 "Chapter02_研究方法": dict(
  arch=[("2.1 跨尺度架構與資料傳遞", "三尺度巢狀概念（Fig4）；Kabsch＋應變檢核；f 縮尺趨勢級定位（防線前置）"),
        ("2.2 門檻滯動潛變模式", "引前作；只寫差異：T 門檻掃描定準法、s1 T=1.0 物理錨點"),
        ("2.3 BPM 襯砌與損傷量化", "G4 鑄造一句；錨定帶；D=斷鍵/registry；三軌分類；QA 一段"),
        ("（2.4 雙向：不設節）", "Ch2 末一句預告，全部內容在 5.1")],
  core=["2.1 核心：傳遞的『一致性』是方法可信度主張——剛體扣除保證只傳變形功、應變檢核 <1e-9。",
        "2.2 核心：T 不是自由參數——乾季近零活化＝定準錨點（掃描法），濕季活化由有效應力路徑自然驅動。",
        "2.3 核心：鍵結斷裂＝微損傷，工程裂縫由三軌法聚合；分母凍結（registry）使 D 可跨情境比較。",
        "【砍】碩論之解析解回顧、模型系譜、建模過程細節皆不入此章。"],
  fig="Fig 4（跨尺度流程＋傳遞）＋Fig 5（三尺度模型三面板）。",
  tab="Table 2（模型組態＋參數＋出處欄）；Table 3（門檻與潛變參數＋物理意義欄）。",
  ref="L2 本構（前作/Sulem/Arora）＋L3（Zheng/Cai 標定）＋（補蒐 BPM 標定 L6）。"),
 "Chapter03_研究案例": dict(
  arch=[("3.1 場址地質與水文", "林鐵沿革一句；南庄層砂頁岩互層＋崩積層；水位計年擺幅與擬合（Fig3）"),
        ("3.2 #46 隧道病徵與監測", "修復履歷時間軸；LiDAR 兩期分級；裂縫計緩剪；三期展開圖（Fig2）；病徵分型"),
        ("3.3 三尺度地質→數值模型", "地層/參數（Table 2 引用）；模型組態（Fig5 引用）；k×100 近壁帶（TT 建議已納）")],
  core=["3.1 核心：水文＝驅動源的證據鏈（水位計實測擺幅→W 面家族→情境放大論證）。",
        "3.2 核心：病徵『分型＋時序』（前作分型手法）——曲線段 957–992 m 為主對象；"
        "109 地震與水文損傷分量並存，為 4.3 對照留伏筆。",
        "3.3 核心：地質模型是數值模型的出處，不是背景裝飾——每參數有孔號/試驗頁碼。"],
  fig="Fig 1（場址綜覽，新繪）＋Fig 2（現地病徵時序）＋Fig 3（水文驅動）。",
  tab="Table 4（現地調查監測資料彙整：LiDAR 兩期/裂縫計/6 孔鑽探/室內試驗）。",
  ref="202411 附錄（App5/6/7/11/13＋附圖93/95/96）＋L5 維養制度文獻。"),
 "Chapter04_數值模擬": dict(
  arch=[("4.1 水文循環下的坡地–隧道尺度響應", "Δpp 檢核；活化區乾衰減/雨爆發/退凍結；Fig6/7/8"),
        ("4.2 襯砌損傷與裂縫發展", "損傷史 A_wet/A_frz（Fig9）→外壓/推力型態（Fig10）→分區與分類＋3D（Fig11/12）"),
        ("4.3 模擬與現地對照", "高損傷帶落點 vs 現地異狀段；同座標展開對照；誠實處理型態差異（D1 防線）")],
  core=["4.1 核心：三尺度互證＝滯動不是單一模型的巧合。段落微結構照前作：圖領句→現象→量化（倍率/區間）→力學詮釋→回扣。",
        "4.2 核心：損傷有節律（相位控制）、有空間（帶狀分區）、有型態（環向主導）。",
        "4.3 核心：對照層級＝分佈與分區，非逐條裂縫；現地含地震/施工縫分量、模擬為水文分量——寫清楚推論範圍。"],
  fig="Fig 6–12 共 7 張（本章主戰場）。",
  tab="（Table 5 之單向欄位）。",
  ref="結果章少引文獻；僅機制詮釋處回扣 P2/P3 文獻。"),
 "Chapter05_討論": dict(
  arch=[("5.1 雙向損傷回饋初探（佔位問句：忽略回饋會低估多少？）", "T5 方法一頁＋Fig13/14＋2.83×＋誠實定位（L0 待分解）"),
        ("5.2 門檻與參數的物理意義（佔位問句：門檻是調出來的嗎？）", "T 定準敏感度；f 縮尺；K0；解析解量級對照（Fahimifar 2009）"),
        ("5.3 工程含義（佔位問句：監測水位能拿來做什麼？）", "滯動泵概念圖（Fig15）；維養決策連結")],
  core=["5.1 核心：回饋使損傷自增強——但誠實區分回饋效應與細時距路徑效應（L0 未跑前用「初探」語級）。",
        "5.2 核心：每個被質疑過的參數（T/f/K0）都有定準邏輯與敏感度說明——口委問題集中殲滅區。",
        "5.3 核心：水位相位＝損傷開關 → 營運隧道的地下水管理即襯砌損傷管理。"],
  fig="Fig 13（雙向方法＋時間線）＋Fig 14（損傷圖演化）＋Fig 15（機理總結，新繪）。",
  tab="Table 5（單向 vs 雙向逐階段對照＋QA gates 摘要）。",
  ref="L3（耦合方法對照）＋Fahimifar 2009＋L1 補強對策（Liu D 2022）。"),
 "Chapter06_結論": dict(
  arch=[("貢獻（條列 5 條）", "①跨尺度鏈量化到單一裂縫級②滯動襯砌尺度證據（91%/7.0/0.0046）"
        "③帶狀分區＋環向主導與現地對照④雙向回饋自增強（2.83×）＋耦合構型論證⑤水位管理＝損傷管理"),
        ("限制（條列 2 條）", "①f 縮尺趨勢級＋100 m 情境放大②雙向為初探（單一回饋律/未含 L0 分解/素混凝土 BPM）")],
  core=["首尾呼應：P5 三層貢獻逐一收口；不復述結果數字堆疊，每條一句精華＋一個代表量。",
        "語級：成功但不自滿——「支持」而非「證明」；限制之後接一句未來展望（L0/敏感度/配筋 BPM）。"],
  fig="無圖。", tab="無表。", ref="無文獻。"),
}

for ch, c in CH.items():
    d = new_doc()
    H(d, f"{ch.split('_')[1]}——子節架構")
    for t, note in c["arch"]:
        P(d, t, bold=True)
        if note:
            P(d, note, size=11, indent=True)
    DEC(d, "子節切法與順序？")
    save(d, ch, "子節架構", "子節架構.docx")

    d = new_doc()
    H(d, f"{ch.split('_')[1]}——段落核心")
    for s in c["core"]:
        P(d, "・" + s)
    save(d, ch, "段落核心", "段落核心.docx")

    d = new_doc()
    H(d, f"{ch.split('_')[1]}——Figure 規劃")
    P(d, c["fig"])
    save(d, ch, "Figure", "圖規劃.docx")

    d = new_doc()
    H(d, f"{ch.split('_')[1]}——Table 規劃")
    P(d, c["tab"])
    save(d, ch, "Table", "表規劃.docx")

    d = new_doc()
    H(d, f"{ch.split('_')[1]}——Reference 掛載")
    P(d, c["ref"])
    save(d, ch, "Reference", "文獻掛載.docx")

print("ALL REVIEW DOCS BUILT")

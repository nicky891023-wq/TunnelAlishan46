#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""build_chapters_v2.py -- Chapter01~06 重新佈署 v2（56 篇精讀經驗灌入；文書規則同規）。

各章五子資料夾（子節架構/段落核心/Figure/Table/Reference）各一 docx。
執行前自動備份各章既有 docx 至該章 _backup_YYMMDD/。
標記語法同 build_ch00_v2：⟦B:藍引用⟧ ⟦R:紅圖⟧ ⟦G:綠表⟧；point= 段末黃底重點。
"""
import io
import shutil
import sys
import time
from pathlib import Path

import docx
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
ROOT = Path(__file__).parent
BLUE, RED, GREEN = (0x1F, 0x4E, 0xC8), (0xC0, 0x00, 0x00), (0x00, 0x80, 0x40)
STAMP = time.strftime("%y%m%d")


def new_doc():
    d = docx.Document()
    st = d.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(12)
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
    return d


def _run(p, text, size, bold=False, color=None, hl=False):
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r.font.bold = bold
    if color:
        r.font.color.rgb = RGBColor(*color)
    if hl:
        r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    rpr = r._r.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        from docx.oxml import OxmlElement
        rf = OxmlElement("w:rFonts")
        rpr.insert(0, rf)
    rf.set(qn("w:ascii"), "Times New Roman")
    rf.set(qn("w:hAnsi"), "Times New Roman")
    rf.set(qn("w:eastAsia"), "標楷體")
    return r


def rich(p, text, size=12, bold=False):
    i = 0
    while i < len(text):
        j = text.find("⟦", i)
        if j < 0:
            _run(p, text[i:], size, bold)
            break
        if j > i:
            _run(p, text[i:j], size, bold)
        k = text.find("⟧", j)
        tag, body = text[j + 1], text[j + 3:k]
        _run(p, body, size, bold, color={"B": BLUE, "R": RED, "G": GREEN}.get(tag))
        i = k + 1


def H(d, text, lv=1):
    p = d.add_paragraph()
    _run(p, text, 15 if lv == 1 else 13, bold=True)
    return p


def P(d, text, size=11.5, bold=False, indent=False, point=None):
    p = d.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Pt(18)
    rich(p, text, size, bold)
    if point:
        _run(p, "　", size)
        _run(p, f"【重點：{point}】", size - 1, color=(0, 0, 0), hl=True)
    return p


def DEC(d, text):
    p = d.add_paragraph()
    _run(p, "【⚖ 待裁決】", 11.5, bold=True, color=RED)
    rich(p, text, 11.5)
    return p


def build(ch, sub, title, body):
    """body = list of ('H'|'P'|'PI'|'D', text[, point]) tuples"""
    d = new_doc()
    H(d, title)
    for item in body:
        kind, text = item[0], item[1]
        pt = item[2] if len(item) > 2 else None
        if kind == "H":
            H(d, text, 2)
        elif kind == "P":
            P(d, text, point=pt)
        elif kind == "PI":
            P(d, text, indent=True, point=pt)
        elif kind == "D":
            DEC(d, text)
    folder = ROOT / ch / sub
    folder.mkdir(parents=True, exist_ok=True)
    name = {"子節架構": "子節架構.docx", "段落核心": "段落核心.docx",
            "Figure": "圖規劃.docx", "Table": "表規劃.docx",
            "Reference": "文獻掛載.docx"}[sub]
    d.save(folder / name)
    print("OK", ch, name)


# 備份
for ch in [p for p in ROOT.iterdir() if p.is_dir() and p.name.startswith("Chapter0") and p.name != "Chapter00_總綱"]:
    bk = ch / f"_backup_{STAMP}"
    bk.mkdir(exist_ok=True)
    for f in ch.rglob("*.docx"):
        if "_backup" not in str(f):
            shutil.copy2(f, bk / f.name.replace(".docx", f"_{f.parent.name}.docx"))
print("chapter backups done")

C1 = "Chapter01_前言"
build(C1, "子節架構", "前言 v2——五段制（B 型雙支線＋二階二分法收束；56 篇統計：眾數 4–5 段）", [
 ("P", "整體原型＝⟦B:(Wang et al., 2024)⟧⟦B:(Chiu et al., 2017)⟧ 的五段雙支線骨架，"
       "P4 內嵌 ⟦B:(Liu et al., 2022)⟧ 二階二分法收束；摘要第 2–3 句先放缺口短版"
       "（⟦B:(Wu et al., 2024)⟧ 缺口說兩次）；各文獻段以缺口句收尾（⟦B:(Paraskevopoulou "
       "& Diederichs, 2018)⟧ 迴聲術：同一缺口全文至少四現）。", "五段=雙支線+二分法收束+缺口迴聲"),
 ("PI", "P1 現象共識＋案例鉤子（工程痛點開場；全段零方法詞）"),
 ("PI", "P2 支線 A：襯砌裂縫/劣化研究 → 第一刀（認識缺口：水文驅動的時間演化未被檢視）"),
 ("PI", "P3 支線 B：水×隧道依時行為 → 第二刀（穩態假設×連續體極限，雙層下刀）"),
 ("PI", "P4 方法系譜＋二階二分法收束（三要素句：水位循環×裂縫顯式解析×跨尺度互制無人同時滿足）"),
 ("PI", "P5 貢獻宣告（To address this gap→案例錨定→方法鏈→not only…but also 雙賣點→機制鏈口號）"),
 ("D", "五段制確認？（4 段裝不下雙支線；6 段嫌鬆）")])
build(C1, "段落核心", "前言各段核心句與句式骨架（落筆作戰卡）", [
 ("H", "P1 核心"),
 ("P", "共識句起手：襯砌裂縫為營運隧道最常見病害（此句即 ⟦B:(Chiu et al., 2017)⟧ 被引 94 次"
       "的主要引用點）→場景收窄（山岳鐵路×多雨×服役 N 年）→硬數字鉤子 1–2 個（裂縫增速/"
       "水位振幅）→維修反覆失效句（⟦B:(Wang et al., 2023)⟧ destruction-restoration 循環）→"
       "需求收尾：掌握水文驅動機制是維養決策前提。", "P1 禁方法詞；鉤子帶數字；失效句殺傷力最高"),
 ("H", "P2 核心"),
 ("P", "一句一人動詞鏈 4–8 筆（觀測→指標→試驗→模擬序）：⟦B:(Chiu et al., 2017)⟧ 演化已記錄"
       "但未檢視乾濕季調製；⟦B:(Wang et al., 2024)⟧ 指標已立但未及生命週期力學根源（其文末"
       "自點名「可由數值模擬驗證」——直接引為墊腳）；⟦B:(Chen et al., 2024)⟧ 靜態裂縫深度"
       "對應；⟦B:(Zhang et al., 2022)⟧ 排水劣化因果。段尾第一刀＝認識缺口型，不砍方法。",
  "第一刀砍認識缺口；方法缺口留 P4"),
 ("H", "P3 核心"),
 ("P", "三軌回顧：解析 ⟦B:(Fahimifar & Zareifard, 2009)⟧（穩態 HM 正典）／實驗 ⟦B:(Sun et "
       "al., 2023)⟧（水位波動響應）⟦B:(Li et al., 2021)⟧（飽和-失水循環劣化）／案例 ⟦B:"
       "(Tarifard et al., 2022)⟧（明言恆定水頭——cyclic vs static 缺口現成）。王牌證詞："
       "⟦B:(Wang et al., 2023)⟧ 自陳水為 root cause 但模型無水。段尾雙短句：穩態假設×"
       "連續體無法顯式解析裂縫（⟦B:(Yin et al., 2022)⟧ 的 FLAC3D 自白可引）。",
  "雙層下刀：情境缺口＋工具缺口分開砍"),
 ("H", "P4 核心"),
 ("P", "系譜句 4–6 筆（⟦B:(Lisjak et al., 2015)⟧⟦B:(Yan et al., 2023)⟧⟦B:(Zhou et al., "
       "2024)⟧⟦B:(Bai et al., 2022)⟧⟦B:(Rasmussen & Min, 2024)⟧）→二階二分法收束：既有"
       "耦合研究聚焦開挖期圍岩；少數及襯砌者僅單調載重；同時滿足（i）水位循環載重（ii）襯砌"
       "裂縫顯式解析（iii）跨尺度互制之框架闕如。搭配 ⟦R:圖4⟧ 系譜面板 This work 落點。",
  "三要素句框出的空格形狀＝本文"),
 ("H", "P5 核心"),
 ("P", "To address this gap→案例錨定句（資料規格式：年數/期數/測點數，⟦B:(Chiu et al., "
       "2017)⟧ P5 寫法）→方法鏈句→not only…but also 雙賣點句（機制揭示＋維養門檻；⟦B:"
       "(Zhang et al., 2025)⟧⟦B:(Ou et al., 2025)⟧ TUST 高頻句式）→機制鏈口號預告（水位循環"
       "→有效應力反覆→裂縫棘輪擴展→性能劣化）。全段不用 novel/first；FDM-DEM 全名首現於此。",
  "not only…but also＝TUST 行情句式；貢獻靠可攜數字自證")])
build(C1, "Figure", "前言圖規劃", [
 ("P", "本章無獨立圖；⟦R:圖4⟧ 之文獻系譜時間軸面板（⟦B:(Zhang et al., 2025)⟧ Fig.1 戰法，"
       "終點紅字 This work）在 P4 引用。可選：P1 現地破壞照片（⟦B:(Kovacevic et al., 2021)⟧ "
       "戰法——視覺證據強化問題真實性）併入 ⟦R:圖2⟧。", "系譜面板=前言的視覺化")])
build(C1, "Table", "前言表規劃", [
 ("P", "⟦G:表1⟧ 文獻案例彙編：8 欄中性矩陣（Author/Case/Lithology/Model/Method/Focus/"
       "Issue/Highlight）約 10–12 列，置於 P2/P3 間；不含 this study 行（前作教訓）；"
       "最後一欄可仿 ⟦B:(Paraskevopoulou & Diederichs, 2018)⟧ 讓表兼當缺口證據"
       "（皆未考慮水位循環）。", "表格代勞文獻回顧＋兼當缺口證據")])
build(C1, "Reference", "前言文獻編隊（掛載對照）", [
 ("P", "P1←L5（⟦B:(Moradi et al., 2021)⟧⟦B:(Ou et al., 2025)⟧）＋⟦B:(Wang et al., 2024)⟧；"
       "P2←支線 A 隊列（Chiu/WangTT/Chen/Sulei/LiuD/Ou）；P3←支線 B 隊列（Fahimifar/Sun/"
       "Li/Tarifard/Wang2023/Yin/Yu）；P4←L3 系譜隊列（Lisjak/Vazaios/Yan/Zhou/Bai/"
       "Rasmussen/Wang2020/Potyondy2015）；P5←無新引。TUST 引文比已 55%，注意補 1–2 篇"
       "經典錨（⟦B:(Potyondy & Cundall, 2004)⟧ 級）做引文雙錨。", "每段的彈藥庫已配置完畢")])

C2 = "Chapter02_研究方法"
build(C2, "子節架構", "研究方法 v2（3 節＋假設前置區；細節引前作與碩論）", [
 ("PI", "2.1 跨尺度架構與資料傳遞（含【建模假設前置區】＝A1 假設條款前置法）"),
 ("PI", "2.2 門檻滯動潛變模式（引前作只寫差異：T 門檻掃描定準法、s1 T=1.0 物理錨點）"),
 ("PI", "2.3 BPM 襯砌與損傷量化（G4 鑄造、錨定帶、D=斷鍵/registry、三軌分類、QA 一段）"),
 ("P", "雙向方法不設節——Ch2 末一句預告、全部內容在 5.1。", "方法章只寫單向；雙向是討論")])
build(C2, "段落核心", "方法各節核心（防禦式寫法）", [
 ("P", "2.1 開頭即設【建模假設清單】：編號條列（趨勢級 f=0.25、水位面家族情境、K0、"
       "邊界處理），每條掛 1–4 篇前例文獻當護盾（⟦B:(Wu et al., 2024)⟧ General "
       "consideration 式；⟦B:(Fahimifar & Zareifard, 2009)⟧ 編號 bullet 式）——先講清楚，"
       "之後全文不再道歉。傳遞一致性（Kabsch＋應變檢核<1e-9）寫成可信度主張。",
  "假設=設定不是懺悔；一次講完不回頭"),
 ("P", "2.2 核心：T 不是自由參數——乾季近零活化＝定準錨點（掃描法）；與前作 λ 的關係一句"
       "交代（尺度與圍壓域不同）。", "門檻定準法=回應口委的正面設計"),
 ("P", "2.3 核心：鍵結斷裂＝微損傷、三軌法聚合成工程裂縫（⟦B:(Zheng et al., 2024)⟧ 同型"
       "定義）；分母凍結（registry 2.08M）使 D 可跨情境比較；QA（CONTROL-0 基線、gates）"
       "一段寫成方法特色。標定敘事引 ⟦B:(Potyondy & Cundall, 2004)⟧⟦B:(Tsang et al., "
       "2023)⟧。", "QA 入正文=審稿人最買單的可信度證據")])
build(C2, "Figure", "方法圖規劃", [
 ("P", "⟦R:圖4⟧ 跨尺度流程＋傳遞＋系譜面板；⟦R:圖5⟧ 三尺度模型三面板。", "兩張撐全章")])
build(C2, "Table", "方法表規劃", [
 ("P", "⟦G:表2⟧ 三尺度模型組態＋參數（含出處欄：孔號/試驗頁碼/文獻）；⟦G:表3⟧ 門檻與"
       "潛變參數（定準依據＋物理意義欄）。", "出處欄=參數來源問題的正面回答")])
build(C2, "Reference", "方法文獻掛載", [
 ("P", "L2 本構（前作/⟦B:(Sulem et al., 1987)⟧）＋L3（Zheng/Bai/Yan）＋L6 標定"
       "（Potyondy/Cho/Yoon/Tsang/Nitka/Peng/Cai）＋假設護盾引文（Wu/Fahimifar）。")])

C3 = "Chapter03_研究案例"
build(C3, "子節架構", "研究案例 v2（案例先行政治：獨立背景節承載證據）", [
 ("PI", "3.1 場址地質與水文（南庄層砂頁岩互層＋崩積層；水位計年擺幅與擬合）"),
 ("PI", "3.2 #46 隧道病徵與監測（修復履歷時間軸→LiDAR 兩期分級→裂縫計緩剪→三期展開圖→病徵分型）"),
 ("PI", "3.3 三尺度地質→數值模型（地層/參數出處；模型組態）"),
 ("P", "戰法＝「案例先行」（⟦B:(Wang et al., 2023)⟧⟦B:(Chiu et al., 2017)⟧⟦B:(Ma et al., "
       "2023)⟧ 同式）：病害調查獨立成節、把證據蒐集前置為客觀背景；前言只留鉤子、細節下沉"
       "此章。", "案例是證據鏈不是背景裝飾")])
build(C3, "段落核心", "案例各節核心（證據三明治佈局）", [
 ("P", "3.1 核心：水文＝驅動源的證據鏈（實測擺幅→擬合→W 面家族情境）；⚖ 主報告水位/雨量"
       "時序仍缺。", "水文證據鏈是 D2 防線的地基"),
 ("P", "3.2 核心：病徵「分型＋時序」——曲線段 957–992 m 為主對象；學 ⟦B:(Tian et al., "
       "2026)⟧ 證據三明治：本章擺實測病徵當問題證據，4.3 再回頭「與現地一致」閉環——"
       "單一資料集賣兩次（動機＋驗證）。震撼數字 hook 提煉 1 個（如裂縫成長率，"
       "⟦B:(Chiu et al., 2017)⟧ 54–87 m/180 天式）。109 地震與水文分量並存＝4.3 對照的"
       "誠實伏筆。", "分型+時序+一個震撼數字；伏筆先埋"),
 ("P", "3.3 核心：每參數有孔號/試驗頁碼（202411 App11/13）；k×100 近壁帶（TT 建議已納）"
       "一句交代。", "地質模型是參數出處")])
build(C3, "Figure", "案例圖規劃", [
 ("P", "⟦R:圖1⟧ 場址綜覽（新繪）；⟦R:圖2⟧ 病徵時序（三期展開圖＋照片＋裂縫計；可含 P1 "
       "破壞照片面板）；⟦R:圖3⟧ 水文驅動。", "圖2 與圖11 同座標=全文最強設計")])
build(C3, "Table", "案例表規劃", [
 ("P", "⟦G:表4⟧ 現地調查監測資料彙整（LiDAR 兩期/裂縫計/6 孔鑽探/室內試驗）＝案例可信度"
       "＝賣點 1 的證據表。", "稀缺性反轉的數據支撐")])
build(C3, "Reference", "案例文獻掛載", [
 ("P", "202411 附錄（App5/6/7/11/13＋附圖93/95/96）＋L5 維養制度（⟦B:(Moradi et al., "
       "2021)⟧）＋分級方法一句。案例資料引用格式⚖：技術報告引用式待定（顧問報告非公開"
       "文獻，可用 unpublished technical report 式或致謝處理——問 TT 慣例）。")])

C4 = "Chapter04_數值模擬"
build(C4, "子節架構", "數值模擬 v2（單向成果 3 節；圖領句微結構貫穿）", [
 ("PI", "4.1 水文循環下的坡地–隧道尺度響應（Δpp 檢核；活化乾衰減/雨爆發/退凍結三尺度互證）"),
 ("PI", "4.2 襯砌損傷與裂縫發展（損傷史→外壓/內力型態→分區分類→三維演化）"),
 ("PI", "4.3 模擬與現地對照（同座標展開；誠實限定框）"),
 ("P", "每結果段落微結構（前作規則）：⟦R:圖X⟧ 領句→現象→量化（倍率/區間）→力學詮釋→"
       "回扣；每節末「綜合前述」收束句。", "圖領句+倍率量化+綜合前述=結果章文法")])
build(C4, "段落核心", "結果各節核心", [
 ("P", "4.1 核心：三尺度互證＝滯動不是單一模型巧合；Δpp=0.98 MPa=ρg·100m 檢核先立可信度。",
  "先檢核後敘事"),
 ("P", "4.2 核心：損傷有節律（A_wet=7.0/A_frz=0.0046/濕窗 91%）、有空間（右肩–右腰帶）、"
       "有型態（環向主導 199/19/3 m）——三個「有」各配一張圖。", "節律/空間/型態三位一體"),
 ("P", "4.3 核心（A9 誠實限定框）：對照層級＝高損傷帶落點與分區、非逐條裂縫；現地含地震/"
       "施工縫分量（109 地震在案）、模擬為水文驅動分量——歸因外部化（⟦B:(Yin et al., "
       "2022)⟧ 量測時機歸因同式）；⟦R:圖2⟧⟦R:圖11⟧ 同座標讓讀者自檢。若 W1 統計出爐，"
       "補一個銳利數字命中（⟦B:(Wu et al., 2024)⟧ 事件級命中戰法：一個 0.282 vs 0.283 "
       "勝過十條大致吻合）。", "誠實+限定+一個銳利命中數字")])
build(C4, "Figure", "結果圖規劃", [
 ("P", "⟦R:圖6⟧ τ-p 門檻｜⟦R:圖7⟧ 活化演化｜⟦R:圖8⟧ 排水暈｜⟦R:圖9⟧ 損傷史｜⟦R:圖10⟧ "
       "外壓＋內力｜⟦R:圖11⟧ 裂縫分區分類（與圖2 同座標）｜⟦R:圖12⟧ 三維演化（hero）。"
       "本章 7 張=全文主戰場。", "7 張各講一件事，不重工")])
build(C4, "Table", "結果表規劃", [
 ("P", "⟦G:表5⟧ 之單向欄位（逐階段 CS-CHK 摘要）。", "表載數字、圖載趨勢")])
build(C4, "Reference", "結果文獻掛載", [
 ("P", "結果章少引文獻；僅機制詮釋處回扣 P2/P3 已引文獻（迴聲不新增）。", "結果章不開新引用戰線")])

C5 = "Chapter05_討論"
build(C5, "子節架構", "討論 v2（3 節；佔位問句法）", [
 ("PI", "5.1 忽略襯砌損傷回饋會低估多少？（雙向 D→E 初探：方法一頁＋2.83×＋誠實定位）"),
 ("PI", "5.2 門檻是調出來的嗎？（T 定準敏感度、f 縮尺、K0；口委問題集中殲滅區）"),
 ("PI", "5.3 監測水位能拿來做什麼？（滯動泵概念圖收尾；維養決策連結）"),
 ("P", "前作技巧：問句當佔位節名、成稿時改正式標題。", "三問句=三個討論事項")])
build(C5, "段落核心", "討論各節核心（包裝落點）", [
 ("P", "5.1 核心：回饋使損傷自增強（2.83×）——但誠實區分回饋效應與細時距路徑效應"
       "（L0 未跑前用「初探」語級＋A3 future work 1:1 鏡像）；力回饋過定性論證"
       "（Picard 發散實證）寫成方法貢獻。", "初探語級；分解留 future work"),
 ("P", "5.2 核心（B1-B6 包裝落點集中地）：f=0.25 用 A4 趨勢級免責框（⟦B:(Barla et al., "
       "2012)⟧ average-response 式）；100m 擺幅用 A2 scope＋needs-based（⟦B:(Tarifard et "
       "al., 2022)⟧ 式）；K0 引文護盾＋一句敏感度；Mitani 解析解對照段：以 ⟦B:(Fahimifar "
       "& Zareifard, 2009)⟧ 做量級 sanity check。", "每個被質疑過的參數都有正面段落"),
 ("P", "5.3 核心：水位相位＝損傷開關→營運隧道地下水管理即襯砌損傷管理；連結分級調查制度"
       "（維養決策者視角收尾＝TUST 甜蜜點）；⟦R:圖15⟧ 滯動泵 bookend（呼應 ⟦R:圖4⟧⟦R:圖6⟧）。",
  "工程含義收尾；概念圖 bookend")])
build(C5, "Figure", "討論圖規劃", [
 ("P", "⟦R:圖13⟧ 雙向方法＋時間線｜⟦R:圖14⟧ D(s,y) 演化｜⟦R:圖15⟧ 滯動泵概念圖（新繪 "
       "pptx，前作 Fig12 手法）。", "討論章 3 張，圖15 收全文")])
build(C5, "Table", "討論表規劃", [
 ("P", "⟦G:表5⟧ 單向 vs 雙向逐階段對照＋QA gates 摘要（審稿防禦表）。", "gates 表=可信度武器")])
build(C5, "Reference", "討論文獻掛載", [
 ("P", "L3 對照（⟦B:(Zhou et al., 2024)⟧ 開挖期 vs 我們營運期差異化）＋⟦B:(Fahimifar & "
       "Zareifard, 2009)⟧ 量級對照＋⟦B:(Liu et al., 2022)⟧ 補強對策一句。")])

C6 = "Chapter06_結論"
build(C6, "子節架構", "結論 v2（條列；無圖無表無文獻）", [
 ("PI", "貢獻 5 條（各一句精華＋一個可攜數字；首尾呼應 P5 三層貢獻）"),
 ("PI", "限制 2 條＋future work 1:1 鏡像（A3 式；只挑 scope 類限制講）"),
 ("P", "語級紀律：支持不證明、無 novel/first、不復述前文數字堆疊（⟦B:(Zhang et al., "
       "2025)⟧ 先揚後抑式收尾）。", "結論=精華不是摘要重播")])
build(C6, "段落核心", "結論條目草案", [
 ("P", "貢獻①跨尺度鏈把水循環→襯砌開裂量化到單一裂縫級②滯動襯砌節律（91%/7.0/0.0046）"
       "③帶狀分區＋環向主導與現地曲線異狀段對照④D→E 回饋自增強（2.83×）＋耦合構型論證"
       "⑤水位相位控制損傷率→地下水管理即損傷管理。", "五條各配一個數字"),
 ("P", "限制①f=0.25 趨勢級＋100m 情境擺幅（→未來：實測年循環重跑）②雙向為初探"
       "（單一回饋律、未含 L0 分解；素混凝土 BPM→未來：配筋 BPM＋對照鏈）。",
  "限制與未來 1:1 鏡像；不提方法論硬傷")])
build(C6, "Figure", "結論圖規劃", [("P", "無圖。")])
build(C6, "Table", "結論表規劃", [("P", "無表。")])
build(C6, "Reference", "結論文獻掛載", [("P", "無文獻（規則）。")])

print("CHAPTERS v2 BUILT")

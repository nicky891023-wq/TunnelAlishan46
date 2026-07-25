#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""build_ch00_v2.py -- Chapter00 總綱（Codex 證據邊界批改版，2026-07-24）。

文書規則（Wade 07-22，IJRMMS 慣例）：
  * 段落末【重點：…】黃底黑字
  * 引用文獻藍字（author-year，如 (Chiu et al., 2017)）
  * 圖XX 紅字、表XX 綠字
行內標記語法：⟦B:藍字⟧ ⟦R:紅字⟧ ⟦G:綠字⟧；P(point=...) 產生段末黃底重點。
執行前自動備份既有 Chapter00 docx 至 _backup_YYMMDD/。
"""
import io
import shutil
import sys
import time
from pathlib import Path

import docx
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
ROOT = Path(__file__).parent
CH0 = ROOT / "Chapter00_總綱"
BLUE, RED, GREEN = (0x1F, 0x4E, 0xC8), (0xC0, 0x00, 0x00), (0x00, 0x80, 0x40)

bk = CH0 / f"_backup_{time.strftime('%y%m%d')}"
bk.mkdir(exist_ok=True)
for f in CH0.glob("*.docx"):
    target = bk / f.name
    if not target.exists():
        shutil.copy2(f, target)
print(f"backup -> {bk.name} ({len(list(bk.glob('*.docx')))} files)")


def new_doc():
    d = docx.Document()
    for sec in d.sections:
        sec.top_margin = Inches(0.85)
        sec.bottom_margin = Inches(0.85)
        sec.left_margin = Inches(0.9)
        sec.right_margin = Inches(0.9)
    st = d.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(12)
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.line_spacing = 1.18
    for name, size in (("Heading 1", 16), ("Heading 2", 14)):
        hs = d.styles[name]
        hs.font.name = "Times New Roman"
        hs.font.size = Pt(size)
        hs.font.bold = True
        hs.font.color.rgb = RGBColor(0, 0, 0)
        hs._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
        hs.paragraph_format.space_before = Pt(10 if name == "Heading 1" else 8)
        hs.paragraph_format.space_after = Pt(5)
        hs.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        ls = d.styles[name]
        ls.font.name = "Times New Roman"
        ls.font.size = Pt(11)
        ls._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
        ls.paragraph_format.space_after = Pt(3)
    return d


def _run(p, text, size, bold=False, color=None, hl=False):
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r.font.bold = bold
    if color:
        r.font.color.rgb = RGBColor(*color)
    if hl:
        r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    rpr = r._r.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.insert(0, rf)
    rf.set(qn("w:ascii"), "Times New Roman")
    rf.set(qn("w:hAnsi"), "Times New Roman")
    rf.set(qn("w:eastAsia"), "標楷體")
    return r


def rich(p, text, size=12, bold=False):
    """render ⟦B:…⟧/⟦R:…⟧/⟦G:…⟧ inline colour markup"""
    i = 0
    while i < len(text):
        j = text.find("⟦", i)
        if j < 0:
            _run(p, text[i:], size, bold)
            break
        if j > i:
            _run(p, text[i:j], size, bold)
        k = text.find("⟧", j)
        tag, body = text[j + 1], text[j + 3:k]
        col = {"B": BLUE, "R": RED, "G": GREEN}.get(tag)
        _run(p, body, size, bold, color=col)
        i = k + 1


def H(d, text, lv=1):
    p = d.add_paragraph(style=f"Heading {lv}")
    _run(p, text, 16 if lv == 1 else 14, bold=True)
    return p


def P(d, text, size=12, bold=False, indent=False, point=None, style=None):
    p = d.add_paragraph(style=style)
    if indent:
        p.paragraph_format.left_indent = Pt(18)
    rich(p, text, size, bold)
    if point:
        _run(p, "　", size)
        _run(p, f"【重點：{point}】", max(size - 1, 10), bold=False, color=(0, 0, 0), hl=True)
    return p


def DEC(d, text):
    p = d.add_paragraph()
    _run(p, "【⚖ 待裁決】", 12, bold=True, color=RED)
    rich(p, text, 12)
    return p


def _cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_dxa(parent, tag, value):
    node = parent.find(qn(tag))
    if node is None:
        node = OxmlElement(tag)
        parent.append(node)
    node.set(qn("w:w"), str(int(value)))
    node.set(qn("w:type"), "dxa")


def _apply_table_geometry(table, weights, section, indent=100):
    """Synchronize tblW/tblGrid/tcW so Word and LibreOffice honor column weights."""
    content = int(section.page_width.twips - section.left_margin.twips - section.right_margin.twips)
    total = content - indent
    weight_sum = float(sum(weights))
    grid = [int(round(total * (w / weight_sum))) for w in weights]
    grid[-1] += total - sum(grid)

    tblPr = table._tbl.tblPr
    _set_dxa(tblPr, "w:tblW", total)
    _set_dxa(tblPr, "w:tblInd", indent)
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tblGrid = table._tbl.tblGrid
    for child in list(tblGrid):
        tblGrid.remove(child)
    for width in grid:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        tblGrid.append(col)

    for row in table.rows:
        for j, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            _set_dxa(tcPr, "w:tcW", grid[j])


def TBL(d, header, rows, size=10.5, widths=None):
    t = d.add_table(rows=1 + len(rows), cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    if widths is None:
        widths = [6.5 / len(header)] * len(header)
    if len(widths) != len(header):
        raise ValueError("table width count must match column count")
    for j, h in enumerate(header):
        c = t.rows[0].cells[j]
        c.text = ""
        c.width = Inches(widths[j])
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _cell_margins(c)
        _run(c.paragraphs[0], h, size + 0.5, bold=True)
        trPr = t.rows[0]._tr.get_or_add_trPr()
        if trPr.find(qn("w:tblHeader")) is None:
            trPr.append(OxmlElement("w:tblHeader"))
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            c = t.rows[1 + i].cells[j]
            c.text = ""
            c.width = Inches(widths[j])
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _cell_margins(c)
            rich(c.paragraphs[0], str(v), size)
        trPr = t.rows[1 + i]._tr.get_or_add_trPr()
        if trPr.find(qn("w:cantSplit")) is None:
            trPr.append(OxmlElement("w:cantSplit"))
    _apply_table_geometry(t, widths, d.sections[0])
    return t


def save(d, name):
    d.save(CH0 / name)
    print("OK", name)


# ========================== 00 題目賣點材料：證據邊界批改版 ==========================
d = new_doc()
H(d, "文章題目、賣點、材料（Codex 證據邊界批改版；2026-07-24）")
P(d, "本頁依 57 篇主庫、現有全文／摘要層級、05 單向定版成果及 06 階段性成果重新校準。"
     "原架構中可用的案例主軸與跨尺度方法予以保留；涉及現地因果、裂縫驗證、雙向回饋與時間尺度者，"
     "改成與現有證據相稱的語氣。", size=11,
  point="本頁是討論稿，不是定版題目；所有主張須能回指 observation、scenario input 或 numerical evidence")

H(d, "一、工作題目（未定版）", 2)
P(d, "現地資料可支持病害存在、多期變化及高異常區分布，但目前不足以單獨證明地下水循環是裂縫的"
     "唯一或直接成因；數值水位變幅亦包含 100 m 放大情境。因此題目宜使用 under、numerical study"
     " 或 numerical investigation，暫不使用 driven by、influenced by、the driving mechanism。",
  point="題目要同時揭露水文情境與數值研究屬性，避免把數值機制寫成已完成的現地因果")
P(d, "建議首選：", bold=True)
P(d, "Lining damage evolution in an operating mountain railway tunnel under cyclic "
     "groundwater-level fluctuations: A cross-scale continuum-discrete numerical study",
  indent=True, bold=True)
P(d, "循環地下水位變動下營運山岳鐵路隧道之襯砌損傷演化：跨尺度連續—離散數值研究",
  indent=True)
P(d, "此題以 damage 包含微損傷與裂縫群聚，以 numerical study 限定證據性質；"
     "「under」只描述研究情境，不預先宣告單因歸因。", indent=True, size=11,
  point="目前最穩健：案例對象＋水文情境＋跨尺度數值範圍")
P(d, "備案 A（裂縫解析證據完成後）：Cross-scale numerical investigation of lining cracking "
     "in an operating mountain railway tunnel under cyclic groundwater-level fluctuations",
  indent=True, size=11)
P(d, "內部備忘（目前不建議作公開候選題目）：Threshold-sensitive time-dependent rock deformation "
     "and lining damage under prescribed groundwater-level cycles。此案仍把尚未獨立驗證的門檻"
     "機制提前升為主軸，且須避免與團隊 in-review 稿件共用 novelty 表述；只有外部已刊文獻與"
     "本文控制分析均能承擔時才重新評估。", indent=True, size=11)
DEC(d, "先裁決文章核心是「案例—機制」或「雙向回饋方法」；在核心未定前不鎖死題目。")

H(d, "二、候選賣點（依目前可守程度排序）", 2)
P(d, "核心 1｜案例導向的工程問題：以營運山岳鐵路隧道之多期 LiDAR／異狀展開圖、裂縫計、"
     "修復履歷、水位與地質調查作為幾何、參數來源及病害分區背景。案例資料用於 case grounding"
     " 與區域層級比較，不宣稱已完成水文因果閉合。", indent=True,
  point="避免「唯一性」；強調資料組合與研究可追溯性")
P(d, "核心 2｜跨尺度整合：把坡地尺度地下水與依時反應、隧道近場傳遞及襯砌 BPM 微損傷放入"
     "同一案例導向數值鏈，並以 Kabsch 剛體扣除與應變一致性檢核控制尺度轉移。前人文獻可分段"
     "支持各環節，但尚不能替本文背書整條因果鏈。", indent=True,
  point="賣點是整合分散機理，不是首次 FLAC–PFC 耦合")
P(d, "方法術語固定分層：05 的尺度串接稱 cross-scale one-way state transfer；PFC–FLAC 同一"
     "模型內的接觸互作稱 continuum–discrete wall-interface coupling；06 的跨交換節點更新稱"
     " externally staggered／partitioned damage–stiffness feedback。三者不可統稱為雙向耦合。",
  indent=True, size=10.5)
P(d, "核心 3｜相位相依的數值損傷節律：在 prescribed 100 m 放大水位情境與既定門檻模型下，"
     "濕窗微損傷事件率與初始乾窗事件率之比 A_wet=7.0，退水後時窗與濕窗事件率之比"
     " A_frz=0.0046。兩者是模型內相位比較指標，不代表現地季節性損傷倍率；「退水後凍結」"
     "只作機制描述，不代表材料完全停止破壞。", indent=True,
  point="可報倍率，但必須附定義、分母、情境與模型適用域")
P(d, "核心 4｜顯式微損傷與空間型態：2.08M 可斷鍵 registry 可解析鍵結破壞事件及其環／斜／縱"
     "群聚；與現地只能比較高異常帶、空間分區與型態語言，不做逐裂縫吻合。單一斷鍵稱 microdamage，"
     "經聚類與幾何重建後才稱工程裂縫。", indent=True,
  point="BPM 斷鍵不是裂縫寬度，也不是未經率定的 RC 服務行為")
P(d, "探索 5｜外部分割式損傷—勁度回饋：06 建立由 PFC 損傷指標更新 FLAC shell 等效勁度的"
     " 26-step 流程。目前 T5 累積斷鍵數與單向 v6 相差 2.83 倍；因兩條流程尚未由同一時間跨度、"
     "同一交換節點設定的 L0 無回饋控制鏈配對，Eeff(D)、交換間隔及 timestep 敏感度亦未完成，"
     "本文不把此差異定義成回饋放大倍率。",
  indent=True, point="06 目前放 Discussion／exploratory；完成控制與敏感度後再決定是否升格")

H(d, "三、材料與證據角色", 2)
TBL(d, ["材料", "在文章中的角色", "現階段可主張／限制"], [
    ["現地病害與履歷", "LiDAR 兩期、異狀展開圖三期、裂縫計與修復紀錄；界定案例及高異常區",
     "可支持病害存在與分區；不可單獨證明地下水成因或逐裂縫驗證"],
    ["現地水文與坡體資料", "界定觀測水位範圍、事件時序與可能外部驅動",
     "約 30 m 現地資訊與 100 m 數值情境必須分欄；缺同步資料時只作背景／邊界條件"],
    ["地質與試驗", "三尺度幾何、地層與參數 provenance；鑽探及室內試驗來源",
     "所有採用值須列來源、轉換與適用尺度；K0、滲透率及門檻另列敏感度"],
    ["05 單向鏈", "文章目前最成熟的主要數值材料；130 天、11 階段、尺度傳遞與襯砌微損傷",
     "可支持相對趨勢與空間機制；f=0.25、100 m 情境與 timestep 限制須前置"],
    ["BPM 率定與 QA", "E、ν、UCS、巴西抗張、破壞型態、起裂門檻、registry 與聚類規則",
     "投稿前應彙成可重現證據；未配筋模型不得稱為完整鋼筋混凝土襯砌"],
    ["06 分割式回饋", "探索開裂—軟化—再分配的正回饋可能性",
     "L0、Eeff(D)、交換間隔與 timestep 對照完成前不作壽命或絕對放大預測"],
    ["57 篇文獻庫", "分段支撐水力、循環路徑、依時變形、坡隧互制、裂縫與耦合方法",
     "A/B 全文可承擔核心論證；B* 只作團隊系譜；C 摘要級僅作背景或待查線索"],
], size=9.3, widths=[1.25, 2.65, 2.8])

H(d, "四、建議的文章策略與待討論事項", 2)
P(d, "現階段建議以「案例—機制型」為主：05 為 Results 主體，06 以 exploratory subsection 放入"
     " Discussion。若 L0、Eeff(D) 與敏感度完成且結果穩健，再考慮把雙向回饋升為核心，或另形成"
     "方法型文章。", point="先讓最成熟的證據承擔主文，避免最弱假設控制整篇審查")
for s in [
    "門檻 T 的物理意義、定準方法與尺度差異",
    "現地水位資訊與 100 m 放大情境的分離及合理用途",
    "f=0.25、PFC timestep scaling 與 trend-level interpretation",
    "現地—模擬比較只到分區／型態／方向，不到單一裂縫",
    "若保留維護意涵，須把斷鍵轉譯成裂縫密度、位置與監測指標，而非直接提出管理門檻",
]:
    P(d, s, size=11, style="List Bullet")
DEC(d, "請 Wade 看完後裁決：①主軸是否採案例—機制型；②題目首選／備案；③06 保留探索性或補證後升格。")
save(d, "00_題目賣點材料.docx")

# ========================== 01 圖表總覽：核心／條件式分層 ==========================
d = new_doc()
H(d, "全文圖表總覽（證據優先版：核心 12 圖＋條件式 3 圖＋5 表）")
P(d, "圖表不先追求數量，而要讓每一項核心主張都有專屬證據。⟦R:圖1–12⟧ 為目前文章主線；"
     "⟦R:圖13–15⟧ 只有在 L0、Eeff(D) 與敏感度完成後才升為正式圖。文獻系譜可濃縮入"
     "⟦R:圖4⟧ 的小面板或前言表，不另占一張主圖。", size=11,
  point="核心圖先閉合案例—輸入—方法—驗證—結果；06 圖不以既有成品自動取得主文席位")
TBL(d, ["圖", "核心訊息", "狀態", "素材／必要加工"], [
    ["⟦R:圖1⟧", "場址、坡體、隧道曲線段與研究尺度的空間關係", "需重繪", "附圖95/96＋App5；統一座標與方向"],
    ["⟦R:圖2⟧", "現地病害證據：分期展開圖、LiDAR、裂縫計與修復事件", "必補 provenance", "每一時期、量測範圍與可比性須標清"],
    ["⟦R:圖3⟧", "現地水文 observation 與 100 m scenario input 分離呈現", "必改", "左：觀測／擬合；右：11-stage 情境及適用域"],
    ["⟦R:圖4⟧", "跨尺度 continuum→continuum→BPM 流程、交換量與證據邊界", "必改", "圖5-01＋5-05；06 回饋以虛線且標 exploratory"],
    ["⟦R:圖5⟧", "三尺度模型幾何、地層、邊界與網格／粒徑", "可整併", "圖5-02/03/04；標出共享區與輸出位置"],
    ["⟦R:圖6⟧", "水力—有效應力—門檻活化的模型內機制", "可用", "圖5-14；補變數定義與時間點"],
    ["⟦R:圖7⟧", "乾季—升水—濕峰—退水的跨尺度活化演化", "可整併", "圖5-08＋5-12；相同色階與三個關鍵時點"],
    ["⟦R:圖8⟧", "襯砌微損傷歷程與 A_wet/A_frz 定義", "可用", "圖5-15；標分母、stage duration 與 100 m scenario"],
    ["⟦R:圖9⟧", "外壓、內力、收斂與損傷位置之對應", "需整併", "圖5-16/17/18；避免只報極值不報位置"],
    ["⟦R:圖10⟧", "BPM 材料率定、CONTROL-0、registry 與聚類 QA", "目前缺主圖", "E/ν/UCS/ft、破壞型態、起裂與守恆／底噪"],
    ["⟦R:圖11⟧", "模擬微損傷群聚與現地病害的同座標分區／型態比較", "必改", "圖5-19＋現地圖；明標不是逐裂縫驗證"],
    ["⟦R:圖12⟧", "三維損傷群聚演化與主要空間帶", "可用", "圖5-20；caption 使用 damage cluster／reconstructed crack"],
    ["⟦R:圖13⟧", "06 分割式交錯回饋流程與 26-tick 狀態傳遞", "條件式", "FIG06-01；不用 strong coupling"],
    ["⟦R:圖14⟧", "同時段同交換節點 L0 vs T5、uniform-D vs mapped-D 與區間敏感度", "尚未具備", "L0、Eeff(D)、2.5/5/10-day 或可行替代對照"],
    ["⟦R:圖15⟧", "可觀測量—數值指標—巡檢分區之工程轉譯", "條件式", "不用「損傷泵」當既定結論；需可執行指標"],
], size=8.7, widths=[0.55, 2.9, 0.85, 2.4])
P(d, "")
TBL(d, ["表", "內容", "必要欄位／用途"], [
    ["⟦G:表1⟧", "文獻證據矩陣", "主題標籤＋方法＋論證角色＋A/B/B*/C；不含 this study 比較列"],
    ["⟦G:表2⟧", "三尺度模型、參數與 provenance", "數值、單位、來源、尺度轉換、使用模型、敏感度狀態"],
    ["⟦G:表3⟧", "門檻、本構、BPM 率定與 QA", "目標值、模擬值、誤差、試件、破壞型態、起裂與聚類門檻"],
    ["⟦G:表4⟧", "現地證據與情境輸入矩陣", "observation／inference／scenario input 分欄，避免證據層級混寫"],
    ["⟦G:表5⟧", "05、L0、T5 與敏感度對照", "相同時間離散、位移、反力、D、斷鍵、分類、QA；L0 未完成前不定版"],
], size=9.2, widths=[0.6, 2.45, 3.65])
DEC(d, "先確認 12 張核心圖是否足以講完 05；06 三圖待控制組完成後再決定是否進主文。")
save(d, "01_圖表總覽.docx")

# ========================== 02 參考文獻主庫：內部證據台帳 ==========================
d = new_doc()
H(d, "參考文獻主庫與證據台帳（內部工作版；59 篇）")
P(d, "主庫目前為 59 篇，其中 TUST 31 篇（53%）；全文精讀 50 篇（85%）。這一份是把書目、蒐集分箱與證據可用性放在一起的"
     "內部工作檔，不是可直接投稿的 APA 7 參考文獻表。正式寫稿時應另輸出純書目清單，移除"
     " [A/B/B*/C]、[TUST] 與分類標頭，並統一作者名、卷期、article number、online-first 狀態及"
     " APA 標點；Crossref 驗證只證明書目存在，不等於全文已核讀。", size=11,
  point="先分清『書目存在』與『內容已核讀』，再決定一篇文獻能承擔多強的句子")
TBL(d, ["等級", "數量", "Access／review 狀態", "暫定 citation readiness"], [
    ["A", "27", "本機出版社 PDF 已直接取得；逐篇核讀狀態另查 note", "引用方法、數據或結論前須回看原文精確頁面"],
    ["B", "22", "FABLE 已以合法全文或本機全文完成全文筆記", "可作核心或支撐文獻；關鍵數字仍回查原文"],
    ["B*", "1", "本團隊 in-review 全文", "只作 research lineage／本文差異，不作唯一外部證據"],
    ["C", "9", "abstract／metadata／題名層級", "只作背景或檢索線索；不得推論章節、圖表及未載數字"],
], size=9.3, widths=[0.65, 0.65, 2.65, 2.75])
P(d, "A† 表示 PDF 已存在但 FABLE 全文筆記尚待同步。正文引用採 Elsevier name-year；"
     "本檔為總庫，逐篇證據仍以 References/reading_notes/ 與原始 PDF 為準。", size=10.5,
  point="核心因果句優先由 A/B 的不同研究團隊交叉支撐；C 不得承擔『前人如何做、結果多少』")
P(d, "分類（v5，Wade 07-25 定義＋子類）：五類串成一條回顧主線——**營運隧道案例與監測 → "
     "地下水及有效應力機制 → 圍岩依時力學模式 → 跨尺度數值傳遞 → 襯砌損傷與維護對策**。"
     "每類再分子類，一個子類＝一個明確的寫作角色（判準＋引用位置），一篇只放一個子類；"
     "分類指出「該文獻在本文中的角色」，仍不得代替其研究對象、荷載路徑與材料之實際涵蓋範圍。",
  size=10.5, point="子類＝寫作角色；證據力另依 A/B/C 分級與逐篇筆記判讀")
P(d, "⚠ 已知缺口：子類 4C（三維工程地質模型建置）目前 0 篇——賣點 3「多尺度三維地質模型」"
     "在前言將無文獻支撐，建議補 2–3 篇（檢索詞：3D engineering geological model／"
     "geological modelling with borehole integration）。", size=10.5, point="缺口已標示，待補")

import json as _json
import re as _re
_acc = _json.loads((ROOT / "References" / "_tools" / "_access_levels.json").read_text(encoding="utf-8"))
_master = (ROOT / "References" / "REFS_MASTER.md").read_text(encoding="utf-8").splitlines()
_doi2id = {}
for _it in _json.loads((ROOT / "References" / "_tools" / "_reading_list.json").read_text(encoding="utf-8")):
    _doi2id[_it["doi"].lower()] = _it["id"]
_direct_pdf = {
    # 本機出版社 PDF 已取得且已完成全文精讀（A 級）
    "WangTT_2010", "Liu_2023", "Bai_2022", "Lionel_2015", "Yuqi_2018",
    "Liu_2019", "WangX_2021", "TianX_2021", "Weixin_2023", "Zhou_2024",
    "Xin_2024", "TianY_2026", "Sulei_2022", "Zheng_2024", "Potyondy_2004",
    "Vazaios_2019", "Lisjak_2015", "Rasmussen_2024", "Wang_2020", "Cho_2007",
    "Yoon_2007", "WangZ_2026", "Bai_2025", "Chang_2024", "Kunjie_2025",
    "Nitka_2018", "Wu_2016",
}
_note_sync_pending = set()
_heading_map = {}
for _ln in _master:
    if _ln.startswith("## "):
        _heading = _ln[3:]
        H(d, _heading_map.get(_heading, _heading), 2)
    elif _ln.startswith("### "):
        _hp = d.add_paragraph()
        _run(_hp, _ln[4:].strip(), 12, bold=True)
        _hp.paragraph_format.left_indent = Pt(8)
        _hp.paragraph_format.space_before = Pt(6)
        _hp.paragraph_format.space_after = Pt(2)
        _hp.paragraph_format.keep_with_next = True
    elif _ln.startswith("- **") or _ln.startswith("- 判準") or _ln.startswith("- 引用位置"):
        _p = P(d, _ln[2:].replace("**", "").strip(), size=9.8, point=None)
        _p.paragraph_format.left_indent = Pt(16)
        _p.paragraph_format.space_after = Pt(1)
    elif _ln.startswith("- ["):
        _m = _re.search(r"10\.[\d.]+/[^\s]+", _ln)
        _id = _doi2id.get((_m.group(0).rstrip(").") if _m else "").lower(), "")
        if "Tsai" in _ln and "in review" in _ln:
            _id = "Tsai_inreview"
        if _id in _direct_pdf:
            _tier = "A†" if _id in _note_sync_pending else "A"
        elif _id == "Tsai_inreview":
            _tier = "B*"
        elif _acc.get(_id, "").startswith("full"):
            _tier = "B"
        else:
            _tier = "C"
        _body = _re.sub(r"^- \[\w\](\s*\*\*\[TUST\]\*\*)?\s*", "", _ln)
        _body = _re.sub(r"^\(被引\s+\d+\)\s*", "", _body)
        _body = _re.sub(r"\s+★.*$", "", _body)
        if _id == "Ma_2023":
            _body = ("Ma, G.; He, Z.; He, C.; Kang, X.; Wang, S.; Xu, G. (2023). "
                     "Time-dependent performance assessment of mountain tunnels considering the hazards "
                     "associated with squeezing and nonuniform steel corrosion of RC lining. "
                     "Computers and Geotechnics, 164, 105808. "
                     "https://doi.org/10.1016/j.compgeo.2023.105808")
        elif _id == "Tsai_inreview":
            _body = ("Tsai, C.-H., Li, H.-H., & Wang, T.-T. (in review). Numerical simulation and "
                     "mechanical interpretation of intermittent time-dependent deformation in tunnels "
                     "[Manuscript under review].")
        _tust = " [TUST]" if "**[TUST]**" in _ln else ""
        _rp = P(d, f"[{_tier}]{_tust} {_body}", size=9.2)
        _rp.paragraph_format.left_indent = Pt(18)
        _rp.paragraph_format.first_line_indent = Pt(-18)
        _rp.paragraph_format.space_after = Pt(2)
        _rp.paragraph_format.line_spacing = 1.02
save(d, "02_參考文獻總集_APA.docx")

# ========================== 03 摘要與關鍵字：05 主體／06 探索 ==========================
d = new_doc()
H(d, "摘要與關鍵字（證據校準工作稿；05 為主、06 為探索）")
P(d, "本版先按「案例—機制型文章」撰寫：摘要以 05 成熟鏈為主，06 僅保留探索性一句。"
     "若後續完成 L0、Eeff(D) 與交換間隔敏感度，再另行升格 06 結果；摘要不以文獻題目統計"
     "取代研究事實，也不放未回填的驗證句。", size=11,
  point="摘要中的每個數字都必須能回指同一版本之圖、表、manifest 或現地資料")
H(d, "建議摘要 A｜案例—機制主線", 2)
P(d, "營運山岳隧道之襯砌病害常與坡體依時變形及水文變動並存，但有限的現地資料通常不足以"
     "分離各成因。既有研究已分別說明地下水載重、水位升降之路徑相依、水—岩依時行為、"
     "坡隧互制與襯砌裂縫演化；然而，循環水位、依時圍岩反應與顯式襯砌損傷如何在案例導向的"
     "跨尺度框架內共同作用，仍未獲充分釐清。本研究以臺灣一座營運山岳鐵路隧道為工程背景，"
     "彙整地質調查、多期 LiDAR、病害展開圖、裂縫計與水位資料。其中地質與水文資料用於模型"
     "及情境建置，巡檢資料則保留作病害分區與型態層級比較；據此建立由坡地尺度、隧道近場至"
     "襯砌 BPM 的三尺度連續—離散數值流程。數值水位循環定位為放大情境，而非實測歷程的直接"
     "重現。130 天、11 階段的單向分析顯示，在本研究情境與門檻"
     "本構下，濕季損傷速率相對初始乾季之比 A_wet 為 7.0，退水後相對濕季之比 A_frz 為"
     " 0.0046，顯示微損傷增量集中於高水位時窗。微損傷事件經空間聚類後，可在分區與型態"
     "層級和現地病害資料對照，但不作逐裂縫驗證。另以 26 個 5-day"
     " tick 進行初步交錯式損傷—勁度更新試算，以檢查流程可行性；現階段不作定量回饋歸因。"
     "研究結果建立了一條可追溯的案例導向跨尺度分析路徑，"
     "可供營運隧道水文監測與病害分區之機制判釋參考。", point="結論限定於模型機制、相對趨勢與分區判釋，不延伸成現地單因證明或壽命預測")
H(d, "摘要禁用／待證明措辭", 2)
for s in [
    "「地下水循環驅動現地裂縫」：目前只能寫 under cyclic groundwater-level fluctuations。",
    "「模擬與現地一致／驗證成功」：目前只允許 spatial-zone／pattern-level comparison。",
    "「2.83 倍由雙向回饋造成」：L0 與敏感度未完成前只能寫 T5 與 v6 的觀察差異。",
    "「RC 襯砌裂縫」：未表現配筋時使用 plain-concrete idealization／BPM lining。",
    "「裂縫數」：原始 bond breaks 應稱 microdamage events；聚類重建後才稱 cracks。",
]:
    P(d, s, size=10.5, style="List Bullet")
H(d, "關鍵字（工作版）", 2)
P(d, "EN: operating tunnel; groundwater-level fluctuation; time-dependent rock deformation; "
     "lining damage; cross-scale modeling; bonded-particle model", size=10.5)
P(d, "中：營運隧道；地下水位變動；岩體依時變形；襯砌損傷；跨尺度模擬；鍵結顆粒模型", size=10.5)
save(d, "03_摘要與關鍵字.docx")

# ========================== 04 碩論濃縮對照表：主張／證據同步 ==========================
d = new_doc()
_sec = d.sections[0]
_sec.orientation = WD_ORIENT.LANDSCAPE
_sec.page_width, _sec.page_height = _sec.page_height, _sec.page_width
_sec.left_margin = _sec.right_margin = Inches(0.65)
_sec.top_margin = _sec.bottom_margin = Inches(0.65)
H(d, "碩論→期刊濃縮對照表（證據優先版）")
P(d, "濃縮不是把碩論章節按比例刪短，而是只保留能服務期刊主問題的證據。現階段建議主問題為："
     "循環地下水情境下，依時坡隧反應如何經跨尺度傳遞形成襯砌微損傷與空間群聚？",
  size=11, point="保留『回答研究問題所需的最少完整證據鏈』，不保留所有曾完成的工作")
H(d, "投稿前不可缺的新增證據", 2)
for s in [
    "現地資料矩陣：每一資料的日期、範圍、精度、座標、是否可跨期比較。",
    "PFC 材料與結構層率定：E、ν、UCS、ft、破壞型態、起裂與尺寸／粒徑敏感度。",
    "尺度傳遞 QA：剛體扣除、應變一致性、映射守恆、CONTROL-0 及版本／hash。",
    "若 06 入核心：同 26-tick L0、Eeff(D)、交換間隔、timestep 與 uniform-D 對照。",
]:
    P(d, s, size=10.5, style="List Bullet")
DEC(d, "請裁決兩項：①05 是否確立為主文核心；②06 是探索性 Discussion，或先補完控制組後再決定篇幅。")
TBL(d, ["來源／內容", "期刊處置／成熟度", "必留材料", "證據邊界／修改"], [
    ["碩論既有｜Ch1 緒論", "重建五段前言｜planning", "案例問題、文獻分段證據、組合缺口、目的",
     "不沿用「無人做過」「完整因果鏈」；57 篇依 A/B/B*/C 選用"],
    ["碩論既有｜Ch2 解析解", "大幅濃縮｜benchmark", "Fahimifar/Tan 作 hydraulic-only benchmark 或量級檢核",
     "不是全砍；若正文沒有實際 benchmark，僅在前言／討論交代其理論位置"],
    ["碩論既有｜Ch2 潛變理論", "只留差異｜verified formulation", "門檻模型、變數、啟閉條件、定準與數值實作",
     "前作 in-review 只作 lineage；外部 Sulem/Paraskevopoulou/Tarifard 承擔必要性"],
    ["碩論既有｜Ch3 案例", "獨立案例節｜field observation", "地質、水文、病害分期、修復與資料 provenance",
     "observation、inference、scenario input 分欄；不把 100 m 寫成現地水位"],
    ["碩論既有｜Ch4 地質模型", "濃縮建模｜verified setup", "三尺度幾何、地層、邊界、參數來源與尺度轉換",
     "刪軟體操作史；保留會影響結果的假設及 K0／滲透率／網格資訊"],
    ["碩論既有＋重分析｜Ch5 方法", "重組方法章｜mixed maturity", "尺度傳遞、Kabsch、門檻、BPM、registry、聚類與 QA",
     "補 PFC E/ν/UCS/ft、起裂、破壞型態與 CONTROL-0；斷鍵不直接等於裂縫"],
    ["碩論定版｜Ch5 單向成果", "Results 主體｜verified numerical result", "水力／門檻反應、損傷節律、載重與空間群聚",
     "過程圖刪；同色階、同時點、同分母；每個倍率皆交代定義與情境"],
    ["碩論既有｜Ch5 侷限", "scope＋uncertainty｜known limitation", "100 m、f=0.25、timestep、素混凝土、單案例",
     "不以「包裝」隱藏；在首次使用前前置，討論末再總結其影響"],
    ["碩論既有｜Ch6 結論", "重寫 3–4 項｜claim audit", "方法整合、相位節律、空間損傷、條件式工程意涵",
     "不把 06 2.83×列核心結論，除非 L0 與敏感度完成"],
    ["畢業後新增｜06 T5", "Discussion exploratory｜preliminary result", "26-tick 流程、D(s,y)→E、觀察差異與 QA",
     "稱 partitioned/staggered feedback；不是 strong coupling；不作絕對壽命預測"],
], size=8.5, widths=[1.15, 1.2, 2.1, 2.25])
save(d, "04_碩論濃縮對照表.docx")

# ========================== 05 寫作工藝手冊：evidence-first ==========================
d = new_doc()
H(d, "寫作工藝手冊（evidence-first 批改版）")
P(d, "本手冊保留 FABLE 對題目、前言與圖領句的有效觀察，但把「缺陷包裝」改為「範圍、假設與"
     "不確定性透明化」。文獻的寫作技巧可以參考，不能取代本研究自己的證據；同刊慣例也不是"
     "隱藏限制或放大 novelty 的理由。", size=11,
  point="寫作目標不是把主張寫強，而是讓主張強度與證據強度完全對齊")

H(d, "一、主張—證據距離", 2)
TBL(d, ["欲寫的句子", "最低所需證據", "本研究目前允許的語氣"], [
    ["現地病害存在／集中", "具日期、座標、範圍與精度的 LiDAR／展開圖／裂縫計", "observed、documented、concentrated in"],
    ["地下水是現地成因", "同步水位—坡移—裂縫時序、替代成因排除或反事實對照", "目前只寫 coexisted with／investigated under，不寫 caused by"],
    ["模型重現現地", "相同指標、座標、尺度、誤差及預先定義判準", "目前寫 zone-/pattern-level comparison，不寫 exact agreement"],
    ["BPM 裂縫", "斷鍵聚類、幾何重建、材料／結構率定及門檻", "原始事件=microdamage；聚類後=reconstructed crack"],
    ["雙向回饋放大", "同一時間跨度與交換節點的 L0、相同起點／輸入、Eeff(D) 與數值敏感度", "目前寫 exploratory difference／may amplify"],
    ["維護門檻／壽命", "多案例或外部驗證、可觀測量、單位一致及決策性能", "目前只寫 implication／mechanistic reference"],
], size=8.8, widths=[1.65, 2.8, 2.25])

H(d, "二、題目與摘要", 2)
P(d, "題目統計可協助判斷可讀性，但不能形成硬規則。「零方法詞」不是目標；當現地因果尚未閉合時，"
     "numerical study／investigation 反而能誠實限定研究屬性。使用 under 描述工況，除非有直接"
     "現地因果證據，避免 driven by、influenced by、driving mechanism。", point="題目先說清楚研究對象、情境與證據類型")
P(d, "摘要依序回答：工程問題→文獻中已知的分段機理→本文要補的組合缺口→案例與方法→三個最重要"
     "結果→適用域。所有數字必須來自同一權威版本；未完成的 W1、L0、敏感度或「待回填」不得"
     "先寫成結果。", point="摘要不使用引用，但每一句在內部都要有可回查來源")

H(d, "三、前言建議脈絡（五段）", 2)
P(d, "P1｜工程問題：營運隧道病害、坡體依時變形與水文波動可能並存；以案例資料規格說明研究"
     "重要性，但不預先指定單一成因。引用營運／維護與裂縫演化文獻。", indent=True, size=11,
  point="先建立問題，不先宣告答案")
P(d, "P2｜水力與循環：由穩態地下水與襯砌水力非均質 ⟦B:(Fahimifar & Zareifard, 2009; "
     "Tan et al., 2018)⟧，推進到升降水路徑相依 ⟦B:(Sun et al., 2023)⟧ 及排水劣化開裂"
     "⟦B:(Zhang et al., 2022)⟧；逐篇明寫其材料、工況與不能外推之處。", indent=True, size=11,
  point="不要把 steady、single-cycle、monotonic blockage 混成同一種 cyclic evidence")
P(d, "P3｜依時坡隧反應：水會改變岩材潛變門檻與速率，坡體依時變形亦會傳至襯砌"
     "⟦B:(Yan et al., 2020; Li et al., 2021; Tarifard et al., 2022; Causse et al., 2015)⟧；"
     "說明材料試驗、概念模型與營運案例的證據尺度不同。", indent=True, size=11,
  point="用文獻接成鏈，但不宣稱任一篇已證明整條鏈")
P(d, "P4｜裂縫與方法：⟦B:Wang (2010)⟧ 支持裂縫型態診斷，⟦B:Chiu et al. (2017)⟧ 支持"
     "時間演化，⟦B:Wang et al. (2024)⟧ 則連結位移與裂縫製圖至剝落評估；三者角色不可互換。BPM／FDEM／"
     "FLAC–PFC 文獻提供顯式破裂與局部離散—外部連續之方法血緣"
     " ⟦B:(Potyondy & Cundall, 2004; Bai et al., 2022; Zhou et al., 2024)⟧，但仍須區分"
     "逐步同步耦合、單體 FDEM、05 單向狀態傳遞與 06 的 5-day 外部分割回饋。", indent=True, size=11,
  point="不可把所有 continuum–discrete 工作都叫作同一種雙向耦合")
P(d, "P5｜組合缺口與目的：限縮為「目前核讀文獻尚未充分釐清循環水位、依時坡隧反應與顯式"
     "襯砌損傷如何在案例導向的跨尺度框架內共同作用」。正式使用 global first／novel 前，"
     "需依確定題目再做系統性"
     "檢索。最後用兩至三個研究問題收束，不先報結果。", indent=True, size=11,
  point="缺口是分散機理的整合，不是首次使用 FLAC、PFC 或兩者耦合")

H(d, "四、Results 與 Discussion 的用語", 2)
TBL(d, ["層級", "Results 寫法", "Discussion 寫法"], [
    ["Observation", "直接報資料、日期、座標、量測值與不確定性", "比較可能機理；不由相關性跳成因果"],
    ["Scenario input", "明列 100 m、stage duration、f、T 與其設定理由", "討論敏感度與適用域；不可稱 field reproduction"],
    ["Numerical response", "報位移、壓力、斷鍵、D、倍率、位置與分母", "以 literature-supported mechanism 解釋，使用 indicates／suggests"],
    ["Field comparison", "先定義比較指標與容許差", "分成位置、型態、方向、量級；不寫逐裂縫吻合"],
    ["06 feedback", "報 L0/T5/U1/L2 及 residual／QA", "只有控制與敏感度閉合後才談 feedback amplification"],
], size=9, widths=[1.25, 2.75, 2.7])

H(d, "五、限制不是包裝，而是可解讀性條件", 2)
P(d, "重要假設在方法首次出現時前置，Discussion 再說明對方向、量級與外推的影響。可不設獨立"
     "Limitations 章，但不能只用一句 future work 帶過。下列六項至少要能被讀者找到：100 m 情境、"
     "f=0.25、PFC timestep scaling、素混凝土／無配筋、單案例、D→E 與 5-day 交換。",
  point="透明揭露會縮小結論範圍，但會提高可信度與可重現性")
P(d, "不得用其他論文的簡化作為本研究的「免責引文」。引文只能說明方法傳統、現象或比較背景；"
     "本研究採用某一假設的理由、影響與敏感度仍須由本研究自己交代。", point="引用不是護盾")

H(d, "六、文獻使用規則", 2)
for s in [
    "A：可直接核對原 PDF；核心數字與方法描述優先使用。",
    "B：全文筆記可支撐主論點，但投稿前對關鍵數字及限定句回查原文。",
    "B*：團隊 in-review 稿件只說明研究 lineage，另配外部已刊文獻。",
    "C：摘要／metadata 只作背景或待取得清單；不得據此寫節序、圖數、limitations 或細部結果。",
    "一個因果主張盡量由不同方法、不同研究團隊的兩篇以上 A/B 文獻交叉支撐。",
]:
    P(d, s, size=10.5, style="List Bullet")

H(d, "七、文書與圖表規則", 2)
P(d, "①段落一事，首句交代角色；②引用 author-year 藍字；③⟦R:圖XX⟧ 紅字、⟦G:表XX⟧ 綠字；"
     "④圖領句後依「讀數→位置／階段→機制→限制」敘述；⑤倍率同時報分子、分母與時間窗；"
     "⑥圖表使用相同座標、色階、單位與版本；⑦黃色【重點】只用於內部審閱，投稿稿移除；"
     "⑧完成正文後做 claim-to-source audit。", point="格式服務論證，不以顏色或修辭代替證據")

H(d, "八、全篇紅線", 2)
for s in [
    "不寫「首次 FLAC–PFC 耦合」或把單體 FDEM 稱為 FLAC–PFC。",
    "不把斷鍵數直接稱裂縫數，不把無配筋 BPM 稱完整 RC 襯砌。",
    "不把 2.83× 全歸因於 feedback，不把 5-day 寫成已收斂的物理時間尺度。",
    "不把 100 m 情境寫成現地紀錄，不把數值趨勢寫成剩餘壽命或維護門檻。",
    "不寫「現地一致」後再留待回填；未完成的比較先標 planned analysis。",
]:
    P(d, s, size=10.5, style="List Bullet")
save(d, "05_寫作工藝手冊.docx")

print("CH00 EVIDENCE-BOUND REVIEW BUILT")

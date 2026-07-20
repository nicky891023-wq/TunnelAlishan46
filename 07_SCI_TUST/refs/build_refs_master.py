#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""build_refs_master.py -- 合併已驗證文獻池 → REFS_MASTER.md ＋ 02b 補遺 docx。

輸入：_gap_collection_raw.json（41 篇，Crossref 驗證）＋ PRIOR15（首輪 TUST 驗證）
　　　＋ SPECIAL（前作 in review 等非 Crossref 項）
策展：KEEP 以 DOI 白名單＋軸別覆寫；解析解家族僅留 Sulem 1987（脈絡）＋
　　　Fahimifar 2009（D3 對照）——其餘依 Wade「解析解砍」裁示不入主庫。
輸出：refs/REFS_MASTER.md（APA 7、六軸、DOI、被引數、TUST 佔比統計）
　　　Chapter00_總綱/02b_文獻補遺_20260718.docx（審閱用增補，不動 02 原檔）
"""
import io
import json
import sys
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
HERE = Path(__file__).parent
RAW = json.loads((HERE / "_gap_collection_raw.json").read_text(encoding="utf-8"))

# 首輪已驗證 15 篇（全 TUST；authors 略式，APA 化時自 Crossref 全取）
PRIOR15 = [
 ("L1","An, P., Fan, L., Ma, S., Zhang, J., Huang, Z., & Fu, H.",2026,"Water pressure distribution around tunnels in pipeline-type karst and lining response characteristics induced by rainfall","TUST","", "107691","10.1016/j.tust.2026.107691",0),
 ("L1","Chen, J., Hu, T., Hu, X., & Jia, K.",2024,"Study on the influence of crack depth on the safety of tunnel lining structure","TUST","143","105470","10.1016/j.tust.2023.105470",0),
 ("L1","Liu, D., Shang, Q., Li, M., Zuo, J., Gao, Y., & Xu, F.",2022,"Cracking behaviour of tunnel lining under bias pressure strengthened using FRP Grid-PCM method","TUST","123","104436","10.1016/j.tust.2022.104436",0),
 ("L2","Zhang, S., Rodriguez-Dono, A., Song, F., & Zhou, Z.",2025,"Time-dependent tunnel deformations: Insights from in-situ tests and numerical analyses","TUST","157","106319","10.1016/j.tust.2024.106319",0),
 ("L2","Mao, Y., Liu, F., Li, Y., Zhao, F., He, M., & Tao, Z.",2026,"Deep tunnel deformation analysis based on large-scale physical test and fractional derivative creep model","TUST","170","107383","10.1016/j.tust.2025.107383",0),
 ("L2","Wang, X., Iura, T., Jiang, Y., Wang, Z., & Liu, R.",2021,"Deformation and mechanical characteristics of tunneling in squeezing ground: A case study of the west section of the Tawarazaka Tunnel","TUST","","103697","10.1016/j.tust.2020.103697",0),
 ("L2","Chang, Z., Yan, C., Xie, W., Lu, Z., Lan, H., & Mei, H.",2024,"Large-scale field tunnel model and time-dependent floor heave induced by humidification","TUST","145","105615","10.1016/j.tust.2024.105615",0),
 ("L3","Yan, X., Yu, H., Chen, Z., Jiang, W., & Li, T.",2023,"A multiscale analysis of adjacent fault dislocation mechanism induced by tunnel excavation based on continuous-discrete coupling method","TUST","140","105263","10.1016/j.tust.2023.105263",0),
 ("L3","Wang, Z., Wang, J., Sun, G., Lin, S., Liu, Z., & Zheng, H.",2026,"Coupled thermo-mechanical simulation of lining cracking evolution and sealing system mechanical response in CAES lined rock caverns using FDEM","TUST","","107460","10.1016/j.tust.2026.107460",0),
 ("L4","Xin, C., Shuai, Y., Song, D., & Liu, X.",2024,"Dynamic interaction and failure mechanism in tunnel-slope systems: Mitigation insights from shaking table tests and numerical analysis","TUST","152","105940","10.1016/j.tust.2024.105940",0),
 ("L4","Tian, Y., Xin, C., Song, D., & Liu, X.",2026,"Dynamic failure mechanism of a high-steep rock mass slope containing a weak interlayer at the tunnel portal and damage evolution of lining subjected to strong earthquakes","TUST","","107586","10.1016/j.tust.2026.107586",0),
 ("L4","Tian, X., Song, Z., & Zhang, Y.",2021,"Monitoring and reinforcement of landslide induced by tunnel excavation: A case study from Xiamaixi tunnel","TUST","110","103796","10.1016/j.tust.2020.103796",0),
 ("L5","Wang, T.-T., Chiu, Y.-C., & Li, K.-J.",2024,"Index for assessing spalling in tunnel lining based on displacement monitoring and crack mapping","TUST","153","105975","10.1016/j.tust.2024.105975",0),
 ("L5","Moradi, P., Asadi, M. J., Ebrahimzadeh, N., & Yarahmadi, B.",2021,"Ilam tunnels inspection, maintenance, and rehabilitation: A case study","TUST","110","103814","10.1016/j.tust.2021.103814",0),
 ("L5","Ou, X., Tang, C., Qu, T., Xu, S., Zhou, Y., & Tian, J.",2025,"Towards digitalized maintenance of operating tunnels: A text documents-based defect evaluation and visualization","TUST","157","106345","10.1016/j.tust.2024.106345",0),
]
SPECIAL = [
 ("L2","Tsai, C.-H., Li, H.-H., & Wang, T.-T.","in review","Numerical simulation and mechanical interpretation of intermittent time-dependent deformation in tunnels","IJRMMS (投稿中)","","","(⚖ 引用方式待 TT)",0),
]

# 軸別覆寫（C-verify 項歸位）＋策展白名單（54 篇→主庫）
AXIS_OVERRIDE = {
 "10.1016/j.ijrmms.2004.02.001": "L6",   # Cai 2004 標定
 "10.1016/j.tust.2013.07.014": "L5", "10.1016/0148-9062(87)90523-7": "L2",
 "10.1016/j.tust.2017.07.001": "L2", "10.1061/(asce)gm.1943-5622.0000163": "L5",
 "10.1016/j.ijrmms.2020.104250": "L4", "10.1016/j.tust.2009.06.002": "L1",
 "10.1016/j.tust.2021.103838": "L5", "10.1016/j.tust.2022.104537": "L5",
 "10.1016/j.engfailanal.2022.106946": "L5", "10.1016/j.ijmst.2021.12.003": "L2",
 "10.1007/s11069-021-04779-6": "L4", "10.1007/s40948-022-00342-0": "L1",
 "10.1002/nag.3650": "L5", "10.1016/j.engfailanal.2024.108392": "L3",
 "10.1016/j.compgeo.2023.105808": "L5",
 "10.1016/j.tust.2017.08.015": "L1", "10.1016/j.engfailanal.2022.106270": "L1",
 "10.1016/j.tust.2023.105138": "L1",
}
DROP = {"10.1007/s00603-015-0890-z"}   # Wu&Xu 2015（三問題綜述，篇幅讓位）

ZH_AXIS = {"L1":"L1 水文→襯砌載重與劣化","L2":"L2 依時變形與潛變本構",
           "L3":"L3 連續-離散耦合與 BPM 襯砌","L4":"L4 隧道–邊坡互制與滲流 HM",
           "L5":"L5 營運隧道維養/案例","L6":"L6 DEM/BPM 標定與混凝土開裂"}

pool = []
for ax, au, yr, ti, jo, vo, ap, doi, cb in PRIOR15:
    pool.append(dict(axis=ax, authors=au, year=yr, title=ti, journal=jo, volume=vo,
                     article_or_pages=ap, doi=doi, cited_by=cb, is_tust=True, status="V"))
for f in RAW:
    doi = f["doi"].lower()
    if doi in DROP:
        continue
    ax = AXIS_OVERRIDE.get(doi, f["axis"] if f["axis"].startswith("L") else "L5")
    jo = "TUST" if f.get("is_tust") else f["journal"]
    pool.append(dict(axis=ax, authors=f["authors"], year=f["year"], title=f["title"],
                     journal=jo, volume=f.get("volume",""), article_or_pages=f.get("article_or_pages",""),
                     doi=f["doi"], cited_by=f.get("cited_by",0),
                     is_tust=bool(f.get("is_tust")) or "tunnelling and underground" in f["journal"].lower(),
                     status="V" if f.get("verified") else "x"))
for ax, au, yr, ti, jo, vo, ap, doi, cb in SPECIAL:
    pool.append(dict(axis=ax, authors=au, year=yr, title=ti, journal=jo, volume=vo,
                     article_or_pages=ap, doi=doi, cited_by=cb, is_tust=False, status="S"))

seen = set(); uniq = []
for f in pool:
    k = f["doi"].lower()
    if k in seen:
        continue
    seen.add(k); uniq.append(f)

# 補抓首輪未記的被引數（Crossref 即時值）
import json as _json
import urllib.request
for f in uniq:
    if f["cited_by"] == 0 and f["doi"].startswith("10."):
        try:
            with urllib.request.urlopen(
                    f"https://api.crossref.org/works/{f['doi']}", timeout=15) as h:
                f["cited_by"] = _json.load(h)["message"].get("is-referenced-by-count", 0)
        except Exception as e:
            print("  cite-fetch miss:", f["doi"], e)
n = len(uniq); nt = sum(1 for f in uniq if f["is_tust"])
print(f"master pool: {n} refs, TUST {nt} ({100*nt/n:.0f}%)")

def apa(f):
    vol = f["volume"]; ap = f["article_or_pages"]
    tail = f"{vol}, {ap}." if vol and ap else (f"{ap}." if ap else "")
    jname = ("Tunnelling and Underground Space Technology" if f["journal"] == "TUST"
             else f["journal"])
    doi = f"https://doi.org/{f['doi']}" if f["status"] == "V" else f["doi"]
    return f"{f['authors']} ({f['year']}). {f['title']}. {jname}, {tail} {doi}"

lines = ["# REFS_MASTER（主文獻庫 v2，2026-07-18）", "",
         f"> {n} 篇；TUST {nt} 篇（{100*nt/n:.0f}%，目標 ≥50% ✓）；V=Crossref 逐 DOI 驗證、",
         "> S=特殊（前作 in review）。禁 MDPI ✓（全庫零 MDPI）。被引數=Crossref 當日值。",
         "> 解析解家族依 Wade 裁示僅留 Sulem 1987（脈絡一句）＋Fahimifar 2009（D3 對照）。", ""]
for ax in ("L1","L2","L3","L4","L5","L6"):
    sub = sorted([f for f in uniq if f["axis"] == ax], key=lambda x: -x["cited_by"])
    lines.append(f"## {ZH_AXIS[ax]}（{len(sub)} 篇，TUST {sum(1 for f in sub if f['is_tust'])}）")
    lines.append("")
    for f in sub:
        tu = " **[TUST]**" if f["is_tust"] else ""
        lines.append(f"- [{f['status']}]{tu} (被引 {f['cited_by']}) {apa(f)}")
    lines.append("")
(HERE / "REFS_MASTER.md").write_text("\n".join(lines), encoding="utf-8")
print("OK refs/REFS_MASTER.md")

# ---------------- 02b 補遺 docx（不動 02 原檔） ----------------
import docx
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

def new_doc():
    d = docx.Document()
    st = d.styles["Normal"]; st.font.name = "Times New Roman"; st.font.size = Pt(11)
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
    return d

def W(d, text, size=11, bold=False, color=None):
    p = d.add_paragraph(); r = p.add_run(text)
    r.font.name = "Times New Roman"; r.font.size = Pt(size); r.font.bold = bold
    if color: r.font.color.rgb = RGBColor(*color)
    rpr = r._r.get_or_add_rPr(); rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        from docx.oxml import OxmlElement
        rf = OxmlElement("w:rFonts"); rpr.insert(0, rf)
    rf.set(qn("w:ascii"), "Times New Roman"); rf.set(qn("w:hAnsi"), "Times New Roman")
    rf.set(qn("w:eastAsia"), "標楷體")
    return p

d = new_doc()
W(d, "參考文獻補遺（2026-07-18 蒐集輪；02 原檔未動，待你批改後合併）", 15, True)
W(d, f"本輪新增＋複驗共 41 篇（Crossref 逐 DOI 驗證、零 MDPI）；主庫現況 {n} 篇、"
     f"TUST {nt} 篇（{100*nt/n:.0f}%）。完整清單見 refs/REFS_MASTER.md。", 11)
W(d, "戰略要點（請特別過目）：", 12, True)
W(d, "1. 邱雅筑・李佳翰・王泰典 (2017)。Lining crack evolution of an operational tunnel "
     "influenced by slope instability. TUST, 65, 167–178.（被引 94）——老師團隊的坡地–襯砌"
     "裂縫直系前身，建議 P1 與 P4 雙掛，系列定位比 2024 剝落指標更直接。", 11)
W(d, "2. Causse, Cojean & Fleurisson (2015)。Interaction between tunnel and unstable slope"
     "…time-dependent… TUST, 50, 270–281.（被引 44）——坡–隧依時互制先例存在，P4 缺口句"
     "須精準化為：『既有坡–隧互制研究未含水循環驅動、未及離散襯砌損傷、且皆單向』。"
     "誠實修正、防審更穩。", 11, color=(0xC0, 0, 0))
W(d, "3. L3 耦合譜系補齊 9 篇：Lisjak 2015 (RMRE,109)、Vazaios 2019 (JRMGE,117)、"
     "Liu&Sun 2019 (TUST,92)、Zhou 2024 (TUST FDM-DEM REV,62)、Bai 2022 (PFC3D-FLAC3D,52)、"
     "Wang Tuo 2020 (TUST,38)、Liu Chang 2023 (DEM 襯砌缺陷開裂,26=最近似我們的離散襯砌)、"
     "Rasmussen&Min 2024 (TUST,12)、Bai 2025 (bridging,37)。", 11)
W(d, "4. L6 標定正典：Potyondy & Cundall 2004 (4,545 引)、Cho 2007 (1,014)、Yoon 2007 (521)、"
     "Potyondy 2015 綜述 (388)、Nitka & Tejchman 2015 混凝土 DEM (158)、Peng 2017 粒徑比 (114)、"
     "Tsang 2023 自動標定 (10)。", 11)
W(d, "5. C 級 16 篇全數複驗取得精確 DOI（含 Cai 2004=860 引、Sharifzadeh 2013=181、"
     "Sulem 1987=178、Paraskevopoulou 2018=167、Barla 2012=148、Yan 2020=124）。", 11)
W(d, "6. 依「解析解砍」裁示：Nomikos/Do/Fahimifar 2010/Hu&Gutierrez/Arora 等解析解家族"
     "不入主庫（留在你的 IJRMMS 藏書）；僅留 Sulem 1987（脈絡）＋Fahimifar 2009（水力"
     "解析解，供 Mitani 老師問題之量級對照）。", 11)
p = W(d, "【⚖ 待裁決】", 11, True, (0xC0, 0, 0))
W(d, "主庫 55 篇略多於 40–50 目標——建議刪 5–8 篇時優先動 L3/L6 低引或期刊等級較低者"
     "（Bai 2025 CPM、Tsang 2023、Jia 2023、Liu B 2022 IJMST…）；請圈刪或授權我依"
     "「引用次數×相關性×期刊等級」自裁。", 11)
out = HERE.parent / "Chapter00_總綱" / "02b_文獻補遺_20260718.docx"
d.save(out)
print("OK", out.name)

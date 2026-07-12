#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
finalize_thesis.py -- post-transplant finishing (run AFTER assemble_thesis + transplant_math):
  pass2: keyword-matched math-paragraph replacements (agent-paraphrased originals)
  pass3: insert the three condensed derivation paragraphs (264/267/271) + Terzaghi table
  data tables: transplant 表3-1 / 表4-1 / 表4-2 from the original docx at their captions
  cleanup: drop duplicate gray placeholders for tables now real
  fonts: enforce Times New Roman + 標楷體 on every w:r run (math m:r untouched)
"""
import copy
import re
from pathlib import Path
import docx
from docx.oxml.ns import qn
from docx.enum.text import WD_COLOR_INDEX

HERE = Path(__file__).parent
SRC = HERE / "260707 碩論_營運中隧道結構受地下水位變化引致圍岩依時變形影響(初稿6)_Ch4-5.docx"
TGT = HERE / "260711_TX碩論_Wade.docx"
M = qn("m:oMath")

def acc(el):
    return "".join((t.text or "") for t in el.iter(qn("w:t")))

src = docx.Document(str(SRC))
tgt = docx.Document(str(TGT))
sp = src.paragraphs

def tag_of(p):
    for r in p.runs[::-1]:
        if r.font.highlight_color is not None and r.text.startswith("("):
            return r.text
    return None

def clone(i):
    el = copy.deepcopy(sp[i]._p)
    for dr in el.iter(qn("w:drawing")):
        dr.getparent().remove(dr)
    return el

# ---------- pass2: keyword replace ----------
MAP = {174: "Terzaghi", 196: "修正Burgers變形模型"}
placed = {}
for i, kw in MAP.items():
    hits = [p for p in tgt.paragraphs
            if kw in p.text and p._p.getparent() is not None
            and not len(list(p._p.iter(M)))]
    if not hits:
        print(f"pass2 MISS {i} ({kw})")
        continue
    p = hits[0]
    t = tag_of(p)
    el = clone(i)
    p._p.addnext(el)
    p._p.getparent().remove(p._p)
    if t:
        r = docx.text.paragraph.Paragraph(el, p._parent).add_run(t)
        r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    placed[i] = el
    print(f"pass2 replaced {i} ({kw}) tag={t}")

# ---------- pass3: insert condensed derivation paragraphs ----------
def find_math_par(prefix):
    for p in tgt.paragraphs:
        if p.text.startswith(prefix) and len(list(p._p.iter(M))):
            return p
    return None

p263 = find_math_par("至於邊界條件")
p268 = find_math_par("在塑性區以及彈性區的交界面上")
for i, anchor, where, tg in ((264, p263, "next", "(公式說明)"),
                             (267, p268, "prev", "(邊界求解)"),
                             (271, p268, "next", "(彈性區推導)")):
    if anchor is None:
        print(f"pass3 MISS anchor for {i}")
        continue
    el = clone(i)
    (anchor._p.addnext if where == "next" else anchor._p.addprevious)(el)
    r = docx.text.paragraph.Paragraph(el, anchor._parent).add_run(tg)
    r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    print(f"pass3 inserted {i} {where} of anchor")

# Terzaghi equation table after para174
if 174 in placed:
    pi = -1
    for child in src.element.body:
        if child.tag == qn("w:p"):
            pi += 1
        elif child.tag == qn("w:tbl") and pi == 173:
            placed[174].addnext(copy.deepcopy(child))
            print("pass3 tbl@173 inserted")
            break

# ---------- data tables 表3-1 / 表4-1 / 表4-2 ----------
# map: src caption prefix -> src table immediately FOLLOWING the caption paragraph
def src_table_after_caption(capkey):
    pi = -1
    want = None
    for child in src.element.body:
        if child.tag == qn("w:p"):
            pi += 1
            t = acc(child).strip().replace(" ", "")
            if re.match(rf"^{capkey}", t):
                want = pi
        elif child.tag == qn("w:tbl") and want is not None:
            return copy.deepcopy(child)
    return None

for capkey, tgtkey in (("表3-?1", "表3-1"), ("表4-?1", "表4-1"), ("表4-?2", "表4-2")):
    tbl = src_table_after_caption(capkey.replace("-?", "-?").replace("表", "表"))
    # src captions may read 表31 / 表 3-1: build regex
    tbl = src_table_after_caption(capkey.replace("-?", "[-]?"))
    if tbl is None:
        print("data table MISS", tgtkey)
        continue
    # target = the placeholder/caption line, NOT a body-text mention
    hit = None
    for p in tgt.paragraphs:
        if p._p.getparent() is None or tgtkey not in p.text:
            continue
        if "佔位" in p.text or "原表保留" in p.text:
            hit = p; break
        if p.text.strip().startswith(tgtkey.replace("表", "表")) and len(p.text) < 60:
            hit = p
    if hit is None:
        print("data table target MISS", tgtkey)
        continue
    hit._p.addnext(tbl)
    newtext = (hit.text.replace("〔插表佔位〕", "").replace("原表保留：", "")
               .replace("〔", "").replace("〕", "").strip())
    for r in hit.runs:
        r.text = ""
    if hit.runs:
        hit.runs[0].text = newtext
    print("data table inserted at caption:", tgtkey)

# ---------- cleanup duplicate gray placeholders ----------
removed = 0
for p in list(tgt.paragraphs):
    if "原表保留：表3-2" in p.text or "原式保留" in p.text:
        p._p.getparent().remove(p._p)
        removed += 1
print("placeholders removed:", removed)

# ---------- corner quotes strip (incl. transplanted originals) ----------
nq = 0
for t in tgt.element.body.iter(qn("w:t")):
    if t.text and ("「" in t.text or "」" in t.text):
        t.text = t.text.replace("「", "").replace("」", "")
        nq += 1
print("corner-quote nodes cleaned:", nq)

# ---------- justify every body paragraph (skip captions/centered) ----------
from docx.enum.text import WD_ALIGN_PARAGRAPH
nj = 0
for p in tgt.paragraphs:
    if (p.style.name or "").startswith("Heading"):
        continue
    if p.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        continue
    if len(p.text.strip()) < 40:      # captions/labels stay as-is
        continue
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    nj += 1
print("justified paragraphs:", nj)

# ---------- final font pass ----------
def enforce_el(root):
    n = 0
    for r in root.iter(qn("w:r")):
        rpr = r.find(qn("w:rPr"))
        if rpr is None:
            rpr = r.makeelement(qn("w:rPr"), {})
            r.insert(0, rpr)
        rf = rpr.find(qn("w:rFonts"))
        if rf is None:
            rf = rpr.makeelement(qn("w:rFonts"), {})
            rpr.insert(0, rf)
        rf.set(qn("w:ascii"), "Times New Roman")
        rf.set(qn("w:hAnsi"), "Times New Roman")
        rf.set(qn("w:eastAsia"), "標楷體")
        n += 1
    return n

n = enforce_el(tgt.element.body)
print("font-enforced runs:", n)

tgt.save(str(TGT))
mp = sum(1 for p in tgt.paragraphs if len(list(p._p.iter(M))))
mt = sum(1 for t in tgt.tables if len(list(t._tbl.iter(M))))
print(f"final: math paras={mp} math tables={mt} paras={len(tgt.paragraphs)} "
      f"images={len(tgt.inline_shapes)} tables={len(tgt.tables)}")

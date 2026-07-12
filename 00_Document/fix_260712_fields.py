#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
fix_260712_fields.py -- merge_260712.py 之後的欄位安全收尾（07-12 深夜定版）。
處理 wt_replace 掃不到的「跨 run 邊界」殘留：
  * 無欄位段落：全文串接後替換、重寫為單一 run（安全）
  * 含欄位段落（REF/SEQ）三處：以明文全文重寫（本文兩段）或以鄰近 Caption 複製保 SEQ（表4-1）
執行順序：python merge_260712.py && python fix_260712_fields.py && Word COM 更新欄位。
"""
import copy
import io
import sys
import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
DST = "260712_TX碩論_Wade.docx"
d = docx.Document(DST)
P = d.paragraphs

def set_text(par, text):
    for r in list(par._p.findall(qn("w:r"))):
        par._p.remove(r)
    run = par.add_run(text)
    run.font.name = "Times New Roman"
    rpr = run._r.get_or_add_rPr()
    rf = OxmlElement("w:rFonts"); rpr.insert(0, rf)
    rf.set(qn("w:ascii"), "Times New Roman")
    rf.set(qn("w:hAnsi"), "Times New Roman")
    rf.set(qn("w:eastAsia"), "標楷體")

def has_field(p):
    return next(p._p.iter(qn("w:fldChar")), None) is not None

PAIRS = [("隧道圍岩擾動尺度", "隧道尺度"), ("隧道圍岩擾動", "隧道尺度"),
         ("物理量僅", "資料僅"), ("圍岩襯砌互制尺度模型", "隧道圍岩-襯砌互制模型"),
         ("圍岩襯砌互制尺度", "隧道圍岩-襯砌互制模型"),
         ("Burger-Mohr", "Burgers-Mohr"), ("2011年", "2021年")]

# 1) 無欄位段落之跨界替換
n = 0
for p in P:
    st = (p.style.name or "").lower()
    if st.startswith("toc") or st == "table of figures" or has_field(p):
        continue
    t = p.text
    t2 = t
    for a, b in PAIRS:
        t2 = t2.replace(a, b)
    if t2 != t:
        set_text(p, t2)
        n += 1
print("cross-boundary plain fixes:", n)

# 2) 含欄位段落之明文重寫（圖4-12 本文、資料傳遞第一段）
def find(key):
    for i, p in enumerate(d.paragraphs):
        if key in p.text:
            return i
    return -1

i = find("該年8月盧碧颱風過境")
if i >= 0 and (has_field(P[i]) or "2011" in P[i].text or "圖 412" in P[i].text):
    set_text(P[i],
     "圖 4-12為案例隧道所在坡地周邊布置的水位計位置，本研究使用2021年自計水壓計測得之年度最高、最低水位資料"
     "(如表 4-1、圖 4-12)擬合後續數值模擬所使用的水位面。該年8月盧碧颱風過境，為當地山區帶來大量豪雨。"
     "觀察各水位計的最高與最低水位，可發現靠近坡腳處水位變化較不明顯，在山區上游地下水位變化則較顯著"
     "(如圖 4-13)，水位升降幅度可達約30 m。")
    print("repaired 圖4-12 body para")
i = find("由大而小的三個模型之間")
if i >= 0 and (has_field(P[i]) or "物理量" in P[i].text or "圖 51" in P[i].text):
    set_text(P[i],
     "由大而小的三個模型之間，資料僅由大尺度單向傳遞至小尺度。初始現地應力由坡地尺度內插至隧道尺度模型，"
     "各水位階段之位移場以邊界速度的形式傳遞給隧道尺度模型，隧道尺度模型中圍岩的位移則傳遞給隧道圍岩-襯砌"
     "互制模型驅動襯砌變形模擬。圖 5-1為資料傳遞之示意，細節於5.2.4節說明。")
    print("repaired 資料傳遞 intro para")

# 3) 表4-1 標題：若含 2011 或已損，以 表4-2 本文 Caption 複製重建（保 SEQ）
def body_caption(prefix):
    for i, p in enumerate(d.paragraphs):
        if (p.style.name or "") == "Caption" and p.text.strip().startswith(prefix):
            for it in p._p.iter(qn("w:instrText")):
                if it.text and "SEQ" in it.text:
                    return i
    return -1

i41 = body_caption("表 4")          # 首個表4 caption（4-1 或已損之 41）
i42 = body_caption("表 4-2")
if i41 >= 0 and i41 != i42 and ("2011" in P[i41].text or "表 41" in P[i41].text):
    cl = copy.deepcopy(P[i42]._p)
    P[i41]._p.addprevious(cl)
    runs = cl.findall(qn("w:r"))
    fe = None
    for k, r in enumerate(runs):
        fc = r.find(qn("w:fldChar"))
        if fc is not None and fc.get(qn("w:fldCharType")) == "end":
            fe = k
    for r in runs[fe + 1:]:
        r.getparent().remove(r)
    capp = Paragraph(cl, P[i41]._parent)
    run = capp.add_run("水位計2021年度最高及最低深度")
    run.font.name = "Times New Roman"
    rpr = run._r.get_or_add_rPr()
    rf = OxmlElement("w:rFonts"); rpr.insert(0, rf)
    rf.set(qn("w:ascii"), "Times New Roman")
    rf.set(qn("w:hAnsi"), "Times New Roman")
    rf.set(qn("w:eastAsia"), "標楷體")
    P[i41]._p.getparent().remove(P[i41]._p)
    print("rebuilt 表4-1 caption with SEQ")

d.save(DST)
print("FIELD-SAFE FIXES SAVED")

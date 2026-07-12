#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
transplant_math.py -- port inline OMML math from 初稿6 into 260710 全文改寫 docx, verbatim.
Wade 2026-07-10: 「Word 數學公式就按照原版貼上不動」.

Method:
  * src math PARAGRAPHS (prose with inline <m:oMath>) replace the matching text-only
    paragraphs in the target (matched by normalized text prefix; drafts copied 原文 so
    prefixes align). The target paragraph's yellow (tag) run is re-appended.
  * src TABLES containing oMath inside the ch2/ch3 body range (equation-number tables and
    表3-2 param table) are inserted right after the target paragraph corresponding to the
    src paragraph immediately preceding each table.
  * gray '原式保留' placeholders in the target are removed afterwards.
"""
import copy
import re
from pathlib import Path
import docx
from docx.oxml.ns import qn

HERE = Path(__file__).parent
SRC = HERE / "260707 碩論_營運中隧道結構受地下水位變化引致圍岩依時變形影響(初稿6)_Ch4-5.docx"
TGT = HERE / "260711_TX碩論_Wade.docx"

M = qn("m:oMath")
W_P = qn("w:p")
W_TBL = qn("w:tbl")

def acc(p_el):
    return "".join((t.text or "") for t in p_el.iter(qn("w:t")))

def norm(s):
    return re.sub(r"[\s，。；：、（）()「」*_>#\[\]]", "", s)

src = docx.Document(str(SRC))
tgt = docx.Document(str(TGT))

# ---- walk src body: paragraph indices + tables with position ----
src_items = []          # (kind, para_index_at_or_before, element)
pi = -1
for child in src.element.body:
    if child.tag == W_P:
        pi += 1
        src_items.append(("p", pi, child))
    elif child.tag == W_TBL:
        src_items.append(("tbl", pi, child))

math_paras = [(i, el) for kind, i, el in src_items
              if kind == "p" and len(list(el.iter(M)))]
# ch2/ch3 region: paragraphs 168..296 (H1 anchors measured earlier)
math_tables = [(i, el) for kind, i, el in src_items
               if kind == "tbl" and len(list(el.iter(M))) and 168 <= i < 296]
print(f"src: {len(math_paras)} math paragraphs, {len(math_tables)} math tables in ch2/ch3")

# ---- target paragraph text index ----
tgt_paras = tgt.paragraphs
tgt_norm = [norm(p.text) for p in tgt_paras]

def find_target(src_text):
    key = norm(src_text)
    for L in (20, 14, 10):
        k = key[:L]
        if len(k) < 6:
            continue
        hits = [j for j, t in enumerate(tgt_norm) if t.startswith(k)]
        if len(hits) == 1:
            return hits[0]
        if len(hits) > 1:
            return hits[0]           # first occurrence (chapters ordered same way)
    # substring fallback on a distinctive mid-slice
    mid = key[10:26]
    if len(mid) >= 8:
        hits = [j for j, t in enumerate(tgt_norm) if mid in t]
        if hits:
            return hits[0]
    return None

def yellow_tag_of(par):
    for r in par.runs[::-1]:
        if r.font.highlight_color is not None and r.text.startswith("("):
            return r.text
    return None

placed = {}
unmatched = []
used_j = set()
for i, el in math_paras:
    text = acc(el)
    j = find_target(text)
    if j is None or j in used_j or tgt_paras[j]._p.getparent() is None:
        unmatched.append((i, text[:30]))
        continue
    used_j.add(j)
    tag = yellow_tag_of(tgt_paras[j])
    new_p = copy.deepcopy(el)
    for dr in new_p.iter(qn("w:drawing")):      # safety: no rel-id carriers
        dr.getparent().remove(dr)
    old_el = tgt_paras[j]._p
    old_el.addnext(new_p)
    parent = old_el.getparent()
    parent.remove(old_el)
    if tag:
        r = docx.text.paragraph.Paragraph(new_p, tgt_paras[j]._parent).add_run(tag)
        r.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW
    placed[i] = new_p
    print(f"  para{i} -> replaced tgt#{j}: {text[:24]!r} tag={tag}")

# rebuild target paragraph cache is unnecessary for tables: we anchor on placed elements
for ti, tel in math_tables:
    # nearest preceding src math paragraph that got placed
    anchor_i = max([i for i in placed if i <= ti], default=None)
    new_t = copy.deepcopy(tel)
    for dr in new_t.iter(qn("w:drawing")):
        dr.getparent().remove(dr)
    if anchor_i is None:
        unmatched.append((f"tbl@{ti}", "no anchor"))
        continue
    placed[anchor_i].addnext(new_t)
    placed[anchor_i] = new_t             # chain subsequent tables after this one
    print(f"  table(after para{ti}) -> inserted after anchor para{anchor_i}")

# ---- drop gray '原式保留' placeholders ----
removed = 0
for p in list(tgt.paragraphs):
    if "原式保留" in p.text:
        p._p.getparent().remove(p._p)
        removed += 1
print("removed placeholders:", removed)
print("UNMATCHED:", unmatched)

tgt.save(str(TGT))
print("saved", TGT.name)

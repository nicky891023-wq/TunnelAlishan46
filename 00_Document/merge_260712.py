#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
merge_260712.py -- 依老師版合稿（07-12）：TT 初稿8 為基底 → 260712_TX碩論_Wade.docx
步驟（冪等，整檔重建）：
  1 複製 TT 原檔
  2 摘要換老師五段版（摘要.docx）＋關鍵字
  3 Ch5 文字回填 v6 定量（雙門檻/黏滯係數/損傷/外壓/推力/變形/裂縫）
  4 換圖 17 張（TT 編號 ↔ result/ v6 圖；同時清該繪圖之 srcRect 裁切、修正長寬）
  5 圖5-21 掛專屬新圖、插入 圖5-22 三維演化＋內文引用
  6 英文摘要改寫對齊
之後由 Word COM 更新欄位。非換圖之既有裁切(srcRect)一律不動。
"""
import copy
import io
import shutil
import sys
import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from PIL import Image

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
SRC = "TTW/260711 碩論_營運中隧道結構受地下水位變化引致圍岩依時變形影響(初稿8)_TT.docx"
DST = "260712_TX碩論_Wade.docx"
R = "result/"
EMU_W = 5270000

shutil.copyfile(SRC, DST)
d = docx.Document(DST)
P = d.paragraphs

def set_text(par, text, en=False):
    for r in list(par.runs):
        r._r.getparent().remove(r._r)
    r = par.add_run(text)
    r.font.name = "Times New Roman"
    rpr = r._r.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.insert(0, rf)
    rf.set(qn("w:ascii"), "Times New Roman")
    rf.set(qn("w:hAnsi"), "Times New Roman")
    if not en:
        rf.set(qn("w:eastAsia"), "標楷體")

def wt_replace(par, pairs):
    """per-w:t substring replace -- preserves fields (SEQ/TOC) and run formatting"""
    changed = False
    for t in par._p.iter(qn("w:t")):
        if t.text:
            new = t.text
            for a, b in pairs:
                new = new.replace(a, b)
            if new != t.text:
                t.text = new
                changed = True
    return changed

# ---------- 2 摘要 ----------
A = docx.Document("TTW/摘要.docx")
abs_all = [p.text.strip() for p in A.paragraphs if p.text.strip()]
kw = [t for t in abs_all if t.startswith("關鍵字")][0]
body = [t for t in abs_all if not t.startswith("關鍵字")]
i27 = [i for i, p in enumerate(P) if p.text.strip() == "摘要"][0] + 1
set_text(P[i27], body[0])
set_text(P[i27 + 1], body[1])
anchor = P[i27 + 1]._p
for t in body[2:]:
    newp = copy.deepcopy(P[i27 + 1]._p)
    anchor.addnext(newp)
    anchor = newp
    set_text(Paragraph(newp, P[i27 + 1]._parent), t)
for p in d.paragraphs:
    if p.text.strip().startswith("關鍵字"):
        set_text(p, kw)
        break
print("step2 abstract ok")

# ---------- 3 文字回填 ----------
def find(key):
    for i, p in enumerate(d.paragraphs):
        if key in p.text:
            return i
    return -1

edits = []
i = find("所有模型統一選用應力門檻閥值T=0.8")
edits.append((i,
 "依時變形參數方面，為了模擬出岩體依時變形啟動範圍於高水位階段增廣、於低水位階段大幅縮減之現象，"
 "水位變化各階段（階段2至11）統一選用應力門檻閥值T=0.8；首階段（初始低水位）另經門檻掃描"
 "（T=0.80至1.00）比較後採T=1.0，以建立乾淨之初始基準，避免初始應力平衡之殘餘擾動誇大低水位期"
 "之依時變形範圍。Maxwell黏滯係數ηm=1.2×10^15 Pa·s，Kelvin黏滯係數ηk=2.4×10^13 Pa·s。"))
i = find("全域統一採用門檻閥值=0.8")
t = d.paragraphs[i].text.replace(
 "全域統一採用門檻閥值=0.8與Maxwell、Kelvin黏滯係數 =1.2×10^15 Pa·s、=2.4×10^13 Pa·s",
 "全域統一採用門檻閥值T=0.8（首階段T=1.0，見5.1節）與Maxwell、Kelvin黏滯係數ηm=1.2×10^15 Pa·s、ηk=2.4×10^13 Pa·s")
edits.append((i, t))
i = find("尖峰38,793條")
edits.append((i,
 "水位變化各階段襯砌損傷演化以PFC模型鍵結斷裂累積情形表示如圖 5-16。扣除首階段之初始基準"
 "（2,751條）後，新增斷鍵數(ΔN)於最高水位(階段6)時達尖峰14,570條，水位循環全程之累積斷鍵密度"
 "約1.70%（含初始基準為1.83%）；升水至最高水位時窗（階段2至6）之斷鍵合計占水位循環總損傷之91%，"
 "顯示大多數微裂隙生成時間集中於水位逐漸升高至最高點的階段，水位降低後新生裂隙數量大減"
 "（末段低水位30天僅新增88條）。若以濕期（階段2至6）平均日斷鍵速率與初始低水位期比較，"
 "比值約7.0；退水後末段低水位期速率僅為濕期平均之0.5%，呈現濕期加速、退水後近乎凍結之滯動特徵"
 "（圖中A_wet、A_frz即此二比值）。"))
i = find("最大約1,832 kPa")
t = d.paragraphs[i].text
t = t.replace("水位最低(階段1及11)及最高(階段時的受壓狀況", "水位最低(階段1及11)及最高(階段6)時的受壓狀況")
t = t.replace("最大約1,832 kPa", "最大約946 kPa")
edits.append((i, t))
i = find("環向內力最大的位置發生在隧道西側拱肩")
t = d.paragraphs[i].text
t = t.replace("顯示環向內力最大的位置發生在隧道西側拱肩，隨水位上升增加、水位下降後不恢復",
 "顯示環向內力最大的位置發生在隧道西側拱肩，最高水位與末段低水位皆約1,030 kN/m（初始低水位同部位約260 kN/m），隨水位上升增加、水位下降後不恢復")
t = t.replace("剖面有向西傾斜之變形，且兩側有內擠",
 "剖面有向西傾斜之變形，且兩側有內擠（最高水位時兩側腰部內擠之區域平均值各約0.15 mm，末段低水位幾乎與其重合）")
edits.append((i, t))
i = find("其中環向裂縫總長度最長")
t = d.paragraphs[i].text
t = t.replace("其中環向裂縫總長度最長，集中於東側腰部，累積總裂縫長度隨水位上升增加",
 "其中環向裂縫總長度最長（末階段環向約199 m，遠大於斜向19 m與縱向3 m；裁除模型兩端各2 m之端部影響帶統計），集中於東側腰部；累積總裂縫長度由初期約6 m於最高水位增至201 m、末階段為221 m，增長集中於水位上升時窗、退水後顯著趨緩")
edits.append((i, t))
for i, t in edits:
    assert i >= 0
    set_text(d.paragraphs[i], t)
print("step3 refill ok:", len(edits))

# ---------- 4 換圖 ----------
P = d.paragraphs
SWAPS = {
 430: ([429], "rId69", R + "圖5-05_物理量傳遞示意.png"),
 439: ([438], "rId70", R + "圖5-02_邊坡尺度模型網格與地下水位面.png"),
 441: ([440], "rId71", R + "圖5-06_邊坡尺度水壓分布.png"),
 449: ([448], "rId72", R + "圖5-03_隧道圍岩擾動尺度模型網格與襯砌.png"),
 451: ([450], "rId73", R + "圖5-11_隧道近場水壓洩降.png"),
 459: ([458], "rId74", R + "圖5-04_圍岩襯砌互制尺度模型.png"),
 486: ([485], "rId79", R + "圖5-08_邊坡尺度塑性區與門檻演化.png"),
 488: ([487], "rId80", R + "圖5-09_邊坡變形歷線.png"),
 495: ([494], "rId81", R + "圖5-10_映射驅動場.png"),
 497: ([496], "rId82", R + "圖5-14_隧道周圍有效應力分布與門檻判定.png"),
 501: ([498, 499, 500], "rId83", R + "圖5-12_隧道圍岩塑性區與門檻演化.png"),
 503: ([502], "rId84", R + "圖5-13_隧道收斂曲線.png"),
 513: ([512], "rId85", R + "圖5-15_襯砌損傷演化歷程.png"),
 515: ([514], "rId86", R + "圖5-16_襯砌外壓展開圖.png"),
 517: ([516], "rId87", R + "圖5-17_襯砌內力分布.png"),
 519: ([518], "rId88", R + "圖5-18_襯砌變形.png"),
 521: ([520], "rId89", R + "圖5-19_襯砌裂縫發展與分類.png"),
}

def fix_drawing(dr, cx, cy):
    for sr in list(dr.iter(qn("a:srcRect"))):
        sr.getparent().remove(sr)
    for ext in dr.iter(qn("wp:extent")):
        ext.set("cx", str(cx)); ext.set("cy", str(cy))
    for ext in dr.iter(qn("a:ext")):
        if ext.getparent().tag == qn("a:xfrm"):
            ext.set("cx", str(cx)); ext.set("cy", str(cy))

rels = d.part.rels
for cap, (imps, rid, png) in SWAPS.items():
    w, h = Image.open(png).size
    cy = int(EMU_W * h / w)
    rels[rid].target_part._blob = open(png, "rb").read()
    first = True
    for ip in imps:
        draws = list(P[ip]._p.iter(qn("w:drawing")))
        for k, dr in enumerate(draws):
            if first and k == 0:
                fix_drawing(dr, EMU_W, cy)
            else:
                dr.getparent().remove(dr)
        first = False
print("step4 swaps ok:", len(SWAPS))

# ---------- 5 圖5-21 專屬圖 + 圖5-22 ----------
png = R + "圖5-21_襯砌裂縫長度累積.png"
w, h = Image.open(png).size
rid21, _ = d.part.get_or_add_image(png)
dr21 = list(P[522]._p.iter(qn("w:drawing")))
assert dr21, "5-21 drawing missing"
for b in dr21[0].iter(qn("a:blip")):
    b.set(qn("r:embed"), rid21)
fix_drawing(dr21[0], EMU_W, int(EMU_W * h / w))
png2 = R + "圖5-20_襯砌裂縫三維演化.png"
w2, h2 = Image.open(png2).size
rid22, _ = d.part.get_or_add_image(png2)
imgp = copy.deepcopy(P[522]._p)
P[523]._p.addnext(imgp)
for b in imgp.iter(qn("a:blip")):
    b.set(qn("r:embed"), rid22)
for dr in imgp.iter(qn("w:drawing")):
    fix_drawing(dr, EMU_W, int(EMU_W * h2 / w2))
capp = copy.deepcopy(P[523]._p)
imgp.addnext(capp)
# keep leading "圖 " + SEQ field runs; drop everything after the field end, append new title
runs = capp.findall(qn("w:r"))
fld_end = None
for k, rr in enumerate(runs):
    fc = rr.find(qn("w:fldChar"))
    if fc is not None and fc.get(qn("w:fldCharType")) == "end":
        fld_end = k
assert fld_end is not None, "5-21 caption has no SEQ field"
for rr in runs[fld_end + 1:]:
    rr.getparent().remove(rr)
cap_par = Paragraph(capp, P[523]._parent)
r = cap_par.add_run(" 隧道襯砌裂縫三維演化圖(四個代表階段)")
r.font.name = "Times New Roman"
rpr = r._r.get_or_add_rPr()
rf = OxmlElement("w:rFonts"); rpr.insert(0, rf)
rf.set(qn("w:ascii"), "Times New Roman"); rf.set(qn("w:hAnsi"), "Times New Roman")
rf.set(qn("w:eastAsia"), "標楷體")
for p in d.paragraphs:
    if "襯砌裂縫展開圖、剖面及各階段長度統計如圖 5-20至圖 5-21" in p.text:
        set_text(p, p.text.replace("如圖 5-20至圖 5-21",
                 "如圖 5-20至圖 5-21（三維視角之四階段演化詳圖 5-22）", 1))
        break
print("step5 fig5-21/5-22 ok")

# ---------- 6 英文摘要 ----------
EN = [
 "During long-term operation, tunnels may be affected by rainfall infiltration, groundwater-level fluctuations, geological structures, and the creep characteristics of the surrounding rock, so that a tunnel structure that has apparently stabilized deforms again, leading to lining cracking, spalling, or support deterioration. Existing studies on tunnel time-dependent deformation focus mostly on the excavation stage or the early post-construction period; how repeated groundwater-level changes during operation modify the effective-stress state of the surrounding rock and further damage the tunnel structure has not been fully discussed. This study therefore takes a mountain railway tunnel in Taiwan as a case to investigate groundwater-level-induced time-dependent deformation of the surrounding rock and its influence on an in-service tunnel structure.",
 "A modified Burgers model with a stress-threshold concept is adopted to establish, on the basis of effective stress and seepage, a conceptual model for the influence of groundwater-level fluctuations on tunnel time-dependent deformation. When the groundwater level rises, pore pressure and seepage forces alter the effective-stress distribution around the tunnel, bringing part of the surrounding rock to approach or reach the time-dependent deformation threshold and thereby initiating or accelerating long-term deformation; when the groundwater level falls, the extent of time-dependent deformation may decrease and the deformation rate slows down.",
 "For the case tunnel, this study integrates existing literature, borehole data, groundwater monitoring, electrical resistivity surveys, terrain interpretation, and tunnel-face images to build a three-dimensional engineering geological model encompassing regional geology and slope investigations through to the tunnel scale, clarifying the spatial relations among tunnel anomalies, crack distribution, strata, weak zones, and groundwater conditions.",
 "To overcome the scale gap between slope-scale hydrogeological processes and local lining damage, a cross-scale numerical simulation framework is established. A slope-scale FLAC3D model first simulates the effective-stress changes and time-dependent slope deformation caused by groundwater-level fluctuations, and transfers the deformation rates to the boundary of the tunnel-scale model. The tunnel-scale model further incorporates excavation, lining, and near-wall pressure drawdown to analyse the stress state, threshold-activated region, and convergence behaviour of the surrounding rock. Finally, a coupled FLAC3D-PFC rock-lining interaction model transfers the rock deformation to the bonded-particle lining to simulate lining compression, damage generation, and crack development.",
 "The results show that groundwater-level fluctuations change the effective stress at the slope scale and, through time-dependent deformation of the surrounding rock, influence the local tunnel response. The slope-scale model exhibits slow deformation at low water levels, accelerated deformation at high levels, and deceleration after drawdown. The rock-lining interaction model reproduces the main characteristics of lining cracking under compression: cracks are dominantly circumferential with subordinate oblique ones and concentrate in the anomalous curved section, in good agreement with field observations. Overall, this study establishes an analysis workflow from groundwater-level change through rock time-dependent deformation to lining damage, and demonstrates that a three-dimensional engineering geological model combined with cross-scale numerical simulation is an effective tool for assessing the long-term deformation mechanism and maintenance strategies of in-service tunnels.",
]
KW = ("Keywords: Lining anomaly, Groundwater-level change, Time-dependent deformation, "
      "3D engineering geological model, Modified Burgers model, Cross-scale numerical simulation")
Pn = d.paragraphs
ei = [i for i, p in enumerate(Pn) if p.text.strip() == "ABSTRACT"][0]
en, kwi = [], None
for i in range(ei + 1, ei + 12):
    t = Pn[i].text.strip()
    if t.startswith("Keywords"):
        kwi = i
        break
    if t:
        en.append(i)
n = min(len(en), len(EN))
for k in range(n):
    set_text(Pn[en[k]], EN[k], en=True)
anchor = Pn[en[-1]]._p
for t in EN[n:]:
    newp = copy.deepcopy(Pn[en[-1]]._p)
    anchor.addnext(newp)
    anchor = newp
    set_text(Paragraph(newp, Pn[en[-1]]._parent), t, en=True)
set_text(Pn[kwi], KW, en=True)
print("step6 EN abstract ok")

# ---------- 7 Codex 合稿審修（NO-GO 清單，07-12 21:5x） ----------
def rep_para(key, new):
    i = find(key)
    assert i >= 0, key[:20]
    set_text(d.paragraphs[i], new)

# blocker: 裂縫排序（TT 原生舊句）
rep_para("圖 5-20結果與現場觀察吻合",
 "圖 5-20至圖 5-22所示結果與現場觀察相符：案例隧道異狀彎道段之襯砌裂縫以環向為主、斜向為輔、"
 "縱向甚少，且集中於特定彎道區段；本研究跨尺度數值模擬重現之裂縫方位序列（環向≫斜向≫縱向）、"
 "集中部位（東側腰部受拉帶）與發生位置（東側彎道段承受坡體推擠處），三者均與現場異狀特徵相容，"
 "符合襯砌開裂係水位升降經圍岩依時變形傳遞至襯砌之推論。")
# major: 5.3.1 s1 舊敘事
rep_para("暫時性造成門檻開啟範圍大",
 "水位升降過程中，坡地尺度模型依時變形門檻開啟範圍變化如圖 5-10。首階段採T=1.0作為初始基準，"
 "開啟數為3,315；其後各階段均採T=0.8，開啟數由階段4之38,374增至階段6（最高水位）之80,165，"
 "退水至階段11降為3,571。首階段與後續階段之門檻係數不同，其數值僅作基準描述；等門檻之階段4、"
 "6、11比較顯示，門檻開啟範圍隨水位上升大幅增加、隨水位下降銳減。")
# major: 7.0 合成註記＋密度分母（重寫損傷段尾）
i = find("尖峰14,570條")
t = d.paragraphs[i].text
t = t.replace("若以濕期（階段2至6）平均日斷鍵速率與初始低水位期比較，"
 "比值約7.0；退水後末段低水位期速率僅為濕期平均之0.5%，呈現濕期加速、退水後近乎凍結之滯動特徵"
 "（圖中A_wet、A_frz即此二比值）。",
 "損傷密度以納入統計之初始可斷鍵結普查總數（約208萬條，扣除拱腳錨定帶等不可斷鍵結）為分母，"
 "與模型全部約223萬條鍵結之口徑不同。若以濕期（階段2至6）平均日斷鍵速率與初始低水位期比較，"
 "比值約7.0；惟首階段採T=1.0、階段2至6採T=0.8，此比值亦包含門檻切換與初始基準性質之差異，"
 "不宜視為地下水位之單一效應。末段低水位與濕期均採T=0.8，末段速率僅為濕期平均之0.5%，"
 "顯示退水後損傷速率明顯降低（圖中A_wet、A_frz即此二比值）。")
set_text(d.paragraphs[i], t)
# major: 5.2.1 「完全由有效應力控制」不實
i = find("使門檻之啟閉完全由有效應力狀態控制")
t = d.paragraphs[i].text.replace(
 "使門檻之啟閉完全由有效應力狀態控制，而非由參數空間差異決定",
 "門檻是否開啟仍取決於各階段有效應力狀態相對於各地層強度包絡之位置")
set_text(d.paragraphs[i], t)
# major: 130 日時程
rep_para("模擬過程共經歷11個階段水位變化",
 "模擬共經歷11個階段水位變化（詳表 5-5），含首階段初始基準合計130日：階段1為初始低水位30日、"
 "階段2至5升水每階5日共20日、階段6最高水位30日、階段7至10退水每階5日共20日、階段11末段低水位"
 "30日。水位面如階梯狀逐階切換並於各階段求解滲流穩態。每一階段中，將坡地尺度模型的節點平均速度，"
 "對應至隧道尺度模型的外邊界，以速度邊界條件推動隧道尺度模型並進行依時變形模擬，示意圖如圖 5-7。")
# major: 資料傳遞增量說明恢復
i = find("並將位移乘以折減係數f=0.25後施加至")
t = d.paragraphs[i].text
t = t.replace("並將位移乘以折減係數f=0.25後施加至隧道圍岩-襯砌互制模型之外邊界。",
 "並將位移乘以折減係數f=0.25後，作為該階段之絕對邊界目標施加至隧道圍岩-襯砌互制模型之外邊界；"
 "互制模型實際載入者為相鄰階段目標之差（增量），以避免累積位移重複施加。")
set_text(d.paragraphs[i], t)
# major: 6.1 傳遞鏈主詞
i = find("坡地尺度模型進一步將隧道尺度模型之階段變形傳遞至")
t = d.paragraphs[i].text.replace(
 "坡地尺度模型進一步將隧道尺度模型之階段變形傳遞至FLAC3D—PFC耦合之隧道圍岩-襯砌互制模型，將岩體變形傳遞至襯砌顆粒鍵結",
 "隧道尺度模型進一步將各階段圍岩變形傳遞至FLAC3D—PFC耦合之隧道圍岩-襯砌互制模型，由岩體變形驅動襯砌顆粒鍵結")
set_text(d.paragraphs[i], t)
# major: 1.3 章名對齊
P2 = d.paragraphs
i13 = find("本文架構主要包含六個章節")
blocks = {"地下水位升降的概念模型": ("研究案例與方法",
  "介紹研究案例隧道環境特性，說明工程地質模型發展方法、考慮地下水滲流效應之岩體水力耦合依時變形模式，並提出地下水位升降與隧道依時變形之概念模型。"),
 "研究案例": ("研究案例三維工程地質模型",
  "依序說明資料蒐集數化與建置流程，以及坡地尺度、隧道尺度三維工程地質模型的調查資料整合與發現。"),
 "數值模擬方法及結果": ("數值模擬方法與結果", None)}
for j in range(i13, i13 + 16):
    tj = P2[j].text.strip()
    if tj in blocks:
        title, desc = blocks[tj]
        set_text(P2[j], title)
        if desc:
            set_text(P2[j + 1], desc)
print("step7a chapter map fixed")
# major: 5.3 驗證措辭 ＋ 近一步
rep_para("藉此驗證案例隧道在特定位置發生裂縫異狀",
 "本節依序呈現坡地尺度、隧道尺度與隧道圍岩-襯砌互制三個模型之模擬結果。模擬的重點是呈現地下水位"
 "變化對於大範圍坡地的有效應力改變，以及應力變化如何進一步對隧道周圍岩體形成推動趨勢，並觀察隧道"
 "在遠場水位升降與周圍滲流雙重作用下，其應力門檻狀態與隧道壁變形行為之反應。應力門檻在低、高水位"
 "階段的開啟情形，以及隧道變形曲線是否呈現低水位趨緩、高水位加速之特徵為觀察焦點，藉此檢視模擬"
 "結果是否與案例隧道特定位置裂縫異狀受地下水位長期影響之假設相容。")
# major: τ-p 路徑措辭
rep_para("應力路徑左移為水位上升時應力變化主要機制",
 "比較最低及最高水位(階段1及6)時，沿著隧道水平方向的應力分布圖，以及隧道圍岩τ-p應力雲圖如圖 5-13。"
 "可見高水位時遠場有效應力降低，τ-p應力點雲整體向較低平均有效應力側移動；本圖呈現不同階段之應力"
 "分布，未追蹤同一單元之完整應力路徑。")
# minor: 圖5-14 塑性區超述
rep_para("使隧道周圍塑性區以及應力作用門檻範圍擴大",
 "水位升降階段隧道尺度模型依時變形門檻作用情形如圖 5-14，亦如同坡地尺度模型，地下水位上升時依時"
 "變形應力門檻開啟範圍擴大，水位恢復後開啟範圍減少。另外，水位抬升時，位於模型主要弱層(具鐵鏽染與"
 "滲水特徵之砂頁岩互層，c=0.10 MPa、φ=30°)的應力點會較早達到應力門檻，使依時變形啟動範圍集中於"
 "隧道所在之弱層範圍內。")
# 4.2 年份（盧碧颱風=2021，TT 之 2011 為筆誤）＋圖4-12 題名重複
for p in d.paragraphs:
    if (p.style.name or "").lower() in ("table of figures",) or (p.style.name or "").lower().startswith("toc"):
        continue
    if "2011年自計水壓計" in p.text:
        wt_replace(p, [("2011年自計水壓計", "2021年自計水壓計")])
    if "水位計2011年度" in p.text:
        wt_replace(p, [("2011年度", "2021年度")])
    if "水位計位置水位計位置" in p.text:
        wt_replace(p, [("水位計位置水位計位置", "水位計位置")])
# 全域術語與校字（跳過 toc 樣式段，目錄由欄位重生）
GLOB = [("隧道圍岩擾動尺度", "隧道尺度"), ("邊坡尺度", "坡地尺度"),
        ("圍岩襯砌互制尺度模型", "隧道圍岩-襯砌互制模型"),
        ("圍岩襯砌互制尺度", "隧道圍岩-襯砌互制模型"),
        ("圍岩—襯砌互制三個模型", "隧道圍岩-襯砌互制三個模型"),
        ("物理量傳遞", "資料傳遞"), ("物理量僅", "資料僅"),
        ("數值模型隧道圍岩擾動需考慮", "隧道尺度數值模型需考慮"),
        ("Burger-Mohr", "Burgers-Mohr"), ("門檻閥值", "門檻係數"),
        ("隧襯襯砌", "隧道襯砌"), ("為例如圖", "為例，如圖"),
        ("四代表階段", "四個代表階段"), ("乾淨之初始基準", "較不受初始擾動影響之基準")]
ng = 0
for p in d.paragraphs:
    st = (p.style.name or "").lower()
    if st.startswith("toc") or st == "table of figures":
        continue
    if wt_replace(p, GLOB):
        ng += 1
print("step7b global terms fixed paras:", ng)

d.save(DST)
print("MERGED SAVED:", DST)

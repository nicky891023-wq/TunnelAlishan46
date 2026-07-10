#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
revise_ch5.py -- apply Wade's Chapter-5 revisions to the 初稿6 Ch4-5 thesis draft as Word
TRACKED CHANGES: (1) terminology unification, (2) expand the coupled 圍岩襯砌互制尺度 results
(currently one sentence) with the crack / loading / damage-evolution findings tied back to the
field crack anomaly, plus (3) framing tie-back sentences. Output -> a NEW file (original kept).
"""
from pathlib import Path
import docx
import tracked

HERE = Path(__file__).parent
SRC  = HERE / "260707 碩論_營運中隧道結構受地下水位變化引致圍岩依時變形影響(初稿6)_Ch4-5.docx"
OUT  = HERE / "260709 碩論_Ch4-5_第五章批改_claude(追蹤修訂).docx"

# ---- terminology map (ordered: specific/longer first; dash variants; drop 隧道 prefix first) ----
TERMS = [
    ("數值模擬方法與結果", "數值方法與結果"),
    ("跨尺度數值模擬策略及方法", "多尺度數值模擬方法"),
    ("跨尺度模型參數設定與邊界條件傳遞", "多尺度模型之參數設定與邊界條件傳遞"),
    ("隧道圍岩—襯砌互制", "圍岩襯砌互制尺度"),
    ("隧道圍岩-襯砌互制", "圍岩襯砌互制尺度"),
    ("隧道圍岩－襯砌互制", "圍岩襯砌互制尺度"),
    ("圍岩—襯砌互制", "圍岩襯砌互制尺度"),
    ("圍岩-襯砌互制", "圍岩襯砌互制尺度"),
    ("圍岩－襯砌互制", "圍岩襯砌互制尺度"),
    ("隧道尺度", "隧道圍岩擾動尺度"),
    ("跨尺度", "多尺度"),
    ("坡地", "邊坡"),
    ("水位靜態", "滲流穩態"),
    ("資料傳遞", "物理量傳遞"),
    ("初始大地應力", "初始現地應力"),
    ("水位變化", "水位升降"),
    ("乾期", "乾季"),
    ("濕期", "雨季"),
    ("初步結果顯示", "數值模擬結果顯示"),
]

def is_content(p):
    nm = p.style.name if p.style else ""
    return not (nm.startswith("toc") or "table of figures" in nm or nm == "Title" or "封面" in nm)

d = docx.Document(str(SRC))

# capture reference paragraphs BEFORE any structural change (objects stay valid across inserts)
def find_para(pred):
    for p in d.paragraphs:
        if pred(p):
            return p
    return None
p_results_intro = None; p_tunnel_res = None; p_coupled_open = None; p_fig512 = None
for i, p in enumerate(d.paragraphs):
    st = p.style.name if p.style else ""
    t = p.text
    if p_results_intro is None and "本節依序呈現" in t: p_results_intro = p
    if p_tunnel_res is None and t.startswith("隨著地下水位由初始低水位"): p_tunnel_res = p
    if p_coupled_open is None and "初步結果顯示" in t: p_coupled_open = p
    # the BODY caption (Normal style, after the coupled opener) -- NOT the table-of-figures entry
    if p_coupled_open is not None and p_fig512 is None and "圖 5-12" in t and "裂縫發展" in t \
            and (p.style.name if p.style else "") == "Normal":
        p_fig512 = p
assert all([p_results_intro, p_tunnel_res, p_coupled_open, p_fig512]), \
    (bool(p_results_intro), bool(p_tunnel_res), bool(p_coupled_open), bool(p_fig512))
body = p_coupled_open.style   # Normal (results body / captions)

# ---- 1. terminology unification (tracked, cross-run robust), doc-wide content paragraphs ----
counts = {}
for p in d.paragraphs:
    if is_content(p):
        for old, n in tracked.tracked_replace_par_terms(p, TERMS).items():
            counts[old] = counts.get(old, 0) + n

# ---- 3. framing tie-back appended to the results intro ----
tracked.append_ins_to_par(
    p_results_intro,
    "本節並以數值所得之襯砌裂縫型態與受力行為，回扣解釋案例隧道異狀彎道段之襯砌開裂異狀。")

# ---- 4. tie-back after the tunnel-scale results (links to the coupled drive) ----
tracked.insert_paragraphs_after(p_tunnel_res, [(
    body,
    "上述隧道圍岩於高水位階段收斂加速、應力門檻作用範圍擴張，低水位階段則趨緩收縮之特性，"
    "即為驅動圍岩襯砌互制尺度模型之邊界條件來源，並進一步造成襯砌之受力與開裂行為。")])

# ---- 5. expand the coupled 圍岩襯砌互制尺度 results (major addition) ----
COUPLED = [
    "進一步以互制模型於各水位階段所累積之斷鍵數量作為襯砌損傷指標，可觀察到襯砌損傷之發展與"
    "水位升降具明確對應關係（圖 5-13）。第一階段之斷鍵屬既有服務歷程損傷，作為基準予以扣除後，"
    "水位循環所誘發之新增損傷於低水位階段發展平緩，於高水位階段（第六階段，最高水位）出現顯著"
    "躍升，單一階段新增斷鍵達三萬八千餘處，其損傷速率約為前期水位上升各階段平均之 1.48 倍；"
    "水位退去後，損傷速率驟降至高水位階段之約 8.5%，呈現「雨季加速、乾季凍結」之依時損傷特徵，"
    "與圍岩依時變形受應力門檻控制之機制一致。至模擬結束，襯砌累積損傷密度約達 4.8%。",

    "就空間分布而言，襯砌損傷主要集中於左右兩側拱腰（腿部）之外緣，並沿隧道軸向延伸成帶狀，"
    "頂拱與拱肩之損傷相對輕微。此一分布與無仰拱馬蹄形斷面於側向擠壓下產生斷面橢圓化、"
    "拱腰承受最大彎矩之力學特性一致。",

    "由於單一斷鍵僅代表顆粒間鍵結之微觀破壞，尚不等同於巨觀裂縫，本研究進一步依裂縫走向對"
    "斷鍵網絡進行分類判別：於各橫斷面（沿軸向每 1 公尺）內連結成弧者判別為環向裂縫，於襯砌"
    "展開面上以局部方向場萃取之連貫斜向條帶判別為斜裂縫，沿軸向延伸者判別為縱向裂縫。分類"
    "結果顯示（圖 5-14），襯砌裂縫以環向裂縫為絕對主體（約占九成以上），斜裂縫居次（主要"
    "分布於拱肩），縱向裂縫則極為稀少。",

    "環向裂縫係拱腰外緣受彎張裂所致，其裂縫面位於橫斷面內、沿環向繞行；於單一斷面呈現一條"
    "沿腿部向下發展之弧形裂縫（圖 5-15），並於沿軸向之各斷面反覆出現、堆疊成帶。位於拱肩之"
    "斜裂縫，則反映該處同時受剪力與軸向作用之複合應力狀態；縱向裂縫之稀少，說明本模型之驅動"
    "係以斷面橢圓化為主導，沿軸向之拉伸彎曲作用相對有限。",

    "就襯砌外壓分布而言，最大外壓出現於拱腳至拱腰一帶（尖峰約 1,900 kPa），與裂縫集中之"
    "位置相互呼應（圖 5-16），顯示拱腰為本斷面受力與開裂之關鍵部位。",

    "上述以環向裂縫為主、斜裂縫居次、幾乎不含軸向裂縫之數值裂縫型態，與案例隧道異狀彎道段"
    "現場所觀察之襯砌裂縫特徵（以環向裂縫為主、斜裂縫為輔）相互吻合。此結果支持本研究之假設："
    "案例隧道特定位置之襯砌開裂異狀，係源自地下水位長期反覆升降所引致之圍岩依時變形，"
    "透過對襯砌之側向擠壓，於拱腰處造成環向受彎張裂。",

    "須說明者，為使離散顆粒襯砌維持於裂縫發展（而非整體崩解）之力學狀態，互制模型所施加之"
    "邊界變形量業經折減，故本節所呈現之損傷量值為趨勢性結果，其空間型態與相對演化行為具"
    "代表性，惟絕對數值不宜逕予外推。",
]
tracked.insert_paragraphs_after(p_coupled_open, [(body, t) for t in COUPLED])

# ---- 6. new figure captions after the existing 圖 5-12 ----
CAPTIONS = [
    "圖 5-13 圍岩襯砌互制尺度模型襯砌損傷演化歷程（各水位階段累積斷鍵數）",
    "圖 5-14 襯砌裂縫展開圖—依裂縫走向分類（環向、斜、縱向）",
    "圖 5-15 襯砌斷面（y=885）裂縫分布—環向弧形裂縫",
    "圖 5-16 襯砌外壓細緻展開圖（拱腰至拱腳外壓最大）",
]
tracked.insert_paragraphs_after(p_fig512, [(body, c) for c in CAPTIONS])

d.save(str(OUT))
print("saved:", OUT.name)
print("terminology replacements:")
for k, v in counts.items():
    if v: print(f"   {k} -> {dict(TERMS)[k]} : {v}")

#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
formal_typeset.py -- 正式論文排版（任務4，07-12）。在 assemble→transplant→finalize 之後執行，
就地改寫 260711_TX碩論_Wade.docx。範本=蔡承翰碩士論文(定稿).docx 實測規格：
  A4、邊界上下2.54/左右3.17cm；Normal=TNR+標楷體12pt、JUSTIFY、1.5倍行距；
  H1=16pt粗體置中、H2/H3=14pt粗體、行距3.0；Caption=10pt（本文取11pt粗體置中沿用現版視覺）。
步驟：
  S1 版面：A4+邊界+標題/內文樣式對齊範本
  S2 封面＋目錄/圖目錄/表目錄（TOC 欄位，Word 開啟後 F9 更新）＋分節與頁碼（前文小寫羅馬/正文阿拉伯）
  S3 圖表標題改掛「圖標題/表標題」樣式（供圖/表目錄欄位抓取）
  S4 式(5-1)(5-2) 灰色佔位→OMML 公式表（式左置中、編號右靠，同移植公式版式）
  S5 內文符號正式化：變數斜體＋上下標（T λ f c φ E ν τ q y z、η_m η_k、A_wet A_frz、
     r_wet r1 r11 rk、E_eq K0 fc ft、σ1′ σ3′ σθ′ p′、u_i v_i d_i R_i Δd_i Δt_i u_{i−1} R_{i−1}、
     ΔN Δh Δpp ux dx dy dz、10^15 10^13）；跳過標題、含 OMML 段、黃底標籤 run
冪等性：S2 以「碩士論文」封面哨兵防重複；S5 以 token 消失自然冪等。
"""
import copy
import re
from pathlib import Path
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm

HERE = Path(__file__).parent
TGT = HERE / "260711_TX碩論_Wade.docx"
M = qn("m:oMath")

doc = docx.Document(str(TGT))

# ============ S1 版面與樣式 ============
for sec in doc.sections:
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.top_margin = sec.bottom_margin = Cm(2.54)
    sec.left_margin = sec.right_margin = Cm(3.17)

st = doc.styles
stn = st["Normal"]
stn.font.size = Pt(12)
stn.paragraph_format.line_spacing = 1.5
stn.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
h1 = st["Heading 1"]; h1.font.size = Pt(16); h1.font.bold = True
h1.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
h1.paragraph_format.line_spacing = 3.0
for nm in ("Heading 2", "Heading 3"):
    h = st[nm]; h.font.size = Pt(14); h.font.bold = True
    h.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
    h.paragraph_format.line_spacing = 3.0
# 圖/表標題樣式（供 TOC \t 抓取）
from docx.enum.style import WD_STYLE_TYPE
for nm in ("圖標題", "表標題"):
    if nm not in [s.name for s in st]:
        cs = st.add_style(nm, WD_STYLE_TYPE.PARAGRAPH)
        cs.base_style = st["Normal"]
        cs.font.size = Pt(11); cs.font.bold = True
        cs.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cs.paragraph_format.line_spacing = 1.5
        rpr = cs.element.get_or_add_rPr()
        rf = rpr.find(qn("w:rFonts"))
        if rf is None:
            rf = OxmlElement("w:rFonts"); rpr.insert(0, rf)
        rf.set(qn("w:ascii"), "Times New Roman"); rf.set(qn("w:hAnsi"), "Times New Roman")
        rf.set(qn("w:eastAsia"), "標楷體")
print("S1 layout+styles done")

# ============ S3 圖表標題掛樣式（先做，S2 依內容插頁不受影響） ============
CAPF = re.compile(r"^圖\s?\d+[-.]\d+")
CAPT = re.compile(r"^表\s?\d+[-.]\d+")
nf = nt = 0
for p in doc.paragraphs:
    t = p.text.strip()
    if len(t) > 80 or not t:
        continue
    if p.alignment == WD_ALIGN_PARAGRAPH.CENTER or True:
        if CAPF.match(t) and len(list(p._p.iter(M))) == 0 and not t.endswith("。"):
            p.style = st["圖標題"]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; nf += 1
        elif CAPT.match(t) and len(list(p._p.iter(M))) == 0 and not t.endswith("。"):
            p.style = st["表標題"]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; nt += 1
print(f"S3 captions styled: fig={nf} tab={nt}")

# ============ S4 式(5-1)(5-2) OMML ============
def mk(tag, txt=None):
    el = OxmlElement(tag)
    if txt is not None:
        el.text = txt
    return el

def mrun(txt, sty=None):
    r = mk("m:r")
    if sty:
        rpr = mk("m:rPr"); s = mk("m:sty"); s.set(qn("m:val"), sty); rpr.append(s); r.append(rpr)
    t = mk("m:t", txt); r.append(t)
    return r

def msub(base, sub, sty="i"):
    s = mk("m:sSub")
    e = mk("m:e"); e.append(mrun(base, sty)); s.append(e)
    sb = mk("m:sub"); sb.append(mrun(sub, "p")); s.append(sb)
    return s

def mfrac(num_els, den_els):
    f = mk("m:f")
    num = mk("m:num"); [num.append(x) for x in num_els]; f.append(num)
    den = mk("m:den"); [den.append(x) for x in den_els]; f.append(den)
    return f

def eq_table(omath_children_groups, number):
    """1x2 borderless table: [equation | (5-x)] following transplanted style."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False
    tbl.columns[0].width = Cm(12.5); tbl.columns[1].width = Cm(2.5)
    c0, c1 = tbl.rows[0].cells
    p0 = c0.paragraphs[0]; p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    om = mk("m:oMath")
    for g in omath_children_groups:
        om.append(g)
    # oMathPara wrapper
    omp = mk("m:oMathPara"); omp.append(om)
    p0._p.append(omp)
    p1 = c1.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p1.add_run(number)
    r.font.size = Pt(12)
    # remove borders
    tpr = tbl._tbl.tblPr
    tb = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{side}"); e.set(qn("w:val"), "none"); tb.append(e)
    tpr.append(tb)
    return tbl

def build_eq51():
    els = []
    els.append(msub("A", "wet"))
    els.append(mrun("="))
    els.append(mfrac([msub("r", "wet")], [msub("r", "1")]))
    els.append(mrun("，　", "p"))
    els.append(msub("r", "wet"))
    els.append(mrun("="))
    els.append(mfrac(
        [mk_group("(ΔN", "2", ")+⋯+(ΔN", "6", ")")],
        [mk_group("(Δt", "2", ")+⋯+(Δt", "6", ")")]))
    return els

def mk_group(pre, s1, mid, s2, post):
    # helper: ΔN_2+⋯+ΔN_6 without literal parens -> build a m:e-like span using d? simpler run chain
    grp = mk("m:e")  # placeholder container not valid at top; instead return list
    return grp

# 簡化：直接以 run+sSub 串接（不需群組容器）
def eq51_els():
    els = [msub("A", "wet"), mrun("="),
           mfrac([msub("r", "wet")], [msub("r", "1")]),
           mrun("，  ", "p"), msub("r", "wet"), mrun("=")]
    num = [msub("ΔN", "2"), mrun("+⋯+"), msub("ΔN", "6")]
    den = [msub("Δt", "2"), mrun("+⋯+"), msub("Δt", "6")]
    els.append(mfrac(num, den))
    return els

def eq52_els():
    return [msub("A", "frz"), mrun("="),
            mfrac([msub("r", "11")], [msub("r", "wet")])]

def replace_placeholder(key, els, number):
    hit = None
    for p in doc.paragraphs:
        if "新式" in p.text and key in p.text:
            hit = p; break
    if hit is None:
        print(f"S4 MISS {key}"); return
    tbl = eq_table(els, number)
    hit._p.addnext(tbl._tbl)
    hit._p.getparent().remove(hit._p)
    print(f"S4 equation {number} placed")

replace_placeholder("A_wet", eq51_els(), "(5-1)")
replace_placeholder("A_frz", eq52_els(), "(5-2)")

# ============ S5 內文符號正式化 ============
# token 處理指令：I=斜體、SUB=(主體斜體,下標正體)、SUP=(主體正體,上標正體)、
# SEQ=多段自訂 [(text, italic?, 'sub'/'sup'/None), ...]
SEQ_TOKens = None
TOKENS = [
    ("10^15", [("10", False, None), ("15", False, "sup")]),
    ("10^13", [("10", False, None), ("13", False, "sup")]),
    ("η_m", [("η", True, None), ("m", False, "sub")]),
    ("η_k", [("η", True, None), ("k", False, "sub")]),
    ("A_wet", [("A", True, None), ("wet", False, "sub")]),
    ("A_frz", [("A", True, None), ("frz", False, "sub")]),
    ("r_wet", [("r", True, None), ("wet", False, "sub")]),
    ("E_eq", [("E", True, None), ("eq", False, "sub")]),
    ("Δd_i", [("Δ", False, None), ("d", True, None), ("i", True, "sub")]),
    ("Δt_i", [("Δ", False, None), ("t", True, None), ("i", True, "sub")]),
    ("u_{i−1}", [("u", True, None), ("i−1", True, "sub")]),
    ("R_{i−1}", [("R", True, None), ("i−1", True, "sub")]),
    ("u_i", [("u", True, None), ("i", True, "sub")]),
    ("v_i", [("v", True, None), ("i", True, "sub")]),
    ("d_i", [("d", True, None), ("i", True, "sub")]),
    ("R_i", [("R", True, None), ("i", True, "sub")]),
    ("σθ′", [("σ", True, None), ("θ", False, "sub"), ("′", False, None)]),
    ("σ1′", [("σ", True, None), ("1", False, "sub"), ("′", False, None)]),
    ("σ3′", [("σ", True, None), ("3", False, "sub"), ("′", False, None)]),
    ("Δpp", [("Δ", False, None), ("p", True, None), ("p", False, "sub")]),
    ("ρgΔh", [("ρ", True, None), ("g", True, None), ("Δ", False, None), ("h", True, None)]),
    ("Δh", [("Δ", False, None), ("h", True, None)]),
    ("ΔNk", [("Δ", False, None), ("N", True, None), ("k", True, "sub")]),
    ("Δtk", [("Δ", False, None), ("t", True, None), ("k", True, "sub")]),
    ("ΔN", [("Δ", False, None), ("N", True, None)]),
    ("−ux", [("−", False, None), ("u", True, None), ("x", True, "sub")]),
    ("dx", [("d", True, None), ("x", True, "sub")]),
    ("dy", [("d", True, None), ("y", True, "sub")]),
    ("dz", [("d", True, None), ("z", True, "sub")]),
    ("K0", [("K", True, None), ("0", False, "sub")]),
    ("fc=", [("f", True, None), ("c", False, "sub"), ("=", False, None)]),
    ("ft=", [("f", True, None), ("t", False, "sub"), ("=", False, None)]),
    ("rk＝", [("r", True, None), ("k", True, "sub"), ("＝", False, None)]),
    ("r1，", [("r", True, None), ("1", False, "sub"), ("，", False, None)]),
    ("r11", [("r", True, None), ("11", False, "sub")]),
    ("p′", [("p", True, None), ("′", False, None)]),
    ("τ-p", [("τ", True, None), ("-", False, None), ("p", True, None)]),
]
# 單字母（獨立出現，非拉丁詞一部分）→ 斜體
SINGLES = "TλfcφEνρτqxyzg"
SINGLE_RE = re.compile(r"(?<![A-Za-z_^{])([TλfcφEνρτqxyzg])(?![A-Za-z_^}])")

def transform_text(text):
    """回傳 [(txt, italic, vert)] 片段清單；無變更回傳 None。"""
    # multi-char tokens 先切
    segs = [(text, None)]  # (str, spec or None)
    for tok, spec in TOKENS:
        out = []
        for s, sp in segs:
            if sp is not None or tok not in s:
                out.append((s, sp)); continue
            parts = s.split(tok)
            for i, part in enumerate(parts):
                if part:
                    out.append((part, None))
                if i < len(parts) - 1:
                    out.append((tok, spec))
        segs = out
    # 單字母斜體
    final = []
    changed = any(sp is not None for _, sp in segs)
    for s, sp in segs:
        if sp is not None:
            for txt, ital, vert in sp:
                final.append((txt, ital, vert))
            continue
        pos = 0
        for m0 in SINGLE_RE.finditer(s):
            if m0.start() > pos:
                final.append((s[pos:m0.start()], False, None))
            final.append((m0.group(1), True, None))
            pos = m0.end(); changed = True
        if pos < len(s):
            final.append((s[pos:], False, None))
    return final if changed else None

def is_yellow(r):
    return r.font.highlight_color is not None

n_sym = 0
for p in doc.paragraphs:
    if (p.style.name or "").startswith("Heading"):
        continue
    if len(list(p._p.iter(M))):
        continue
    for r in list(p.runs):
        if is_yellow(r) or not r.text:
            continue
        pieces = transform_text(r.text)
        if pieces is None:
            continue
        anchor = r._r
        for txt, ital, vert in pieces:
            nr = copy.deepcopy(anchor)
            # clear text, set new
            for t in nr.findall(qn("w:t")):
                nr.remove(t)
            t = OxmlElement("w:t"); t.text = txt
            t.set(qn("xml:space"), "preserve")
            nr.append(t)
            rpr = nr.find(qn("w:rPr"))
            if rpr is None:
                rpr = OxmlElement("w:rPr"); nr.insert(0, rpr)
            # italic
            for tag in ("w:i", "w:iCs", "w:vertAlign"):
                e = rpr.find(qn(tag))
                if e is not None:
                    rpr.remove(e)
            if ital:
                rpr.append(OxmlElement("w:i")); rpr.append(OxmlElement("w:iCs"))
            if vert:
                va = OxmlElement("w:vertAlign")
                va.set(qn("w:val"), "superscript" if vert == "sup" else "subscript")
                rpr.append(va)
            anchor.addprevious(nr)
        anchor.getparent().remove(anchor)
        n_sym += 1
print(f"S5 symbol runs transformed: {n_sym}")

# ============ S2 封面＋目錄＋分節頁碼（最後做，插於最前） ============
def para_before(anchor_p, text="", size=12, bold=False, center=True, style=None):
    newp = OxmlElement("w:p")
    anchor_p._p.addprevious(newp)
    from docx.text.paragraph import Paragraph
    P = Paragraph(newp, anchor_p._parent)
    if style:
        P.style = st[style]
    if text:
        r = P.add_run(text)
        r.font.size = Pt(size); r.font.bold = bold
        r.font.name = "Times New Roman"
        rpr = r._r.get_or_add_rPr()
        rf = rpr.find(qn("w:rFonts"))
        if rf is None:
            rf = OxmlElement("w:rFonts"); rpr.insert(0, rf)
        rf.set(qn("w:eastAsia"), "標楷體")
    P.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    return P

def field_para(anchor_p, instr):
    P = para_before(anchor_p, "", center=False)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instr)
    r = OxmlElement("w:r"); t = OxmlElement("w:t")
    t.text = "（在 Word 中按 F9 更新目錄）"
    r.append(t); fld.append(r)
    P._p.append(fld)
    return P

def page_break_before(anchor_p):
    P = para_before(anchor_p, "", center=False)
    r = P.add_run(); br = OxmlElement("w:br"); br.set(qn("w:type"), "page")
    r._r.append(br)
    return P

first = doc.paragraphs[0]
if "碩士論文" not in "".join(q.text for q in doc.paragraphs[:12]):
    COVER = [
        ("國立臺灣大學工學院土木工程學系", 18, True),
        ("碩士論文", 18, True),
        ("Department of Civil Engineering", 14, False),
        ("College of Engineering", 14, False),
        ("National Taiwan University", 14, False),
        ("Master Thesis", 14, False),
        ("", 12, False),
        ("營運中隧道結構受地下水位變化引致圍岩依時變形影響", 16, True),
        ("Influence of Time-Dependent Deformation of Surrounding Rock Induced by "
         "Groundwater-Level Fluctuations on In-Service Tunnel Structures", 14, False),
        ("", 12, False),
        ("郭婷軒", 16, True),
        ("Ting-Hsuan Kuo", 14, False),
        ("", 12, False),
        ("指導教授：王泰典 博士", 16, True),
        ("Advisor: Tai-Tien Wang, Ph.D.", 14, False),
        ("", 12, False),
        ("中華民國115年7月", 16, True),
        ("July 2026", 14, False),
    ]
    for text, size, bold in COVER:
        para_before(first, text, size=size, bold=bold, center=True)
    page_break_before(first)
    print("S2 cover inserted")
else:
    print("S2 cover already present")

# 目錄三頁：插在第一章之前（摘要之後）
ch1p = None
for p in doc.paragraphs:
    if p.text.strip().startswith("第一章") and (p.style.name or "").startswith("Heading"):
        ch1p = p; break
if ch1p is not None and "目　錄" not in "".join(q.text for q in doc.paragraphs[:80]):
    para_before(ch1p, "目　錄", 16, True, True)
    field_para(ch1p, r'TOC \o "1-3" \h \z \u')
    page_break_before(ch1p)
    para_before(ch1p, "圖目錄", 16, True, True)
    field_para(ch1p, r'TOC \h \z \t "圖標題,1"')
    page_break_before(ch1p)
    para_before(ch1p, "表目錄", 16, True, True)
    field_para(ch1p, r'TOC \h \z \t "表標題,1"')
    page_break_before(ch1p)
    print("S2 TOC/LOF/LOT inserted")

# 分節：第一章前結束前文節（羅馬頁碼），正文節阿拉伯重編
def set_footer_pagenum(sec, fmt=None, start=None, roman=False):
    sec.footer.is_linked_to_previous = False
    fp = sec.footer.paragraphs[0]
    for r in list(fp.runs):
        r._r.getparent().remove(r._r)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
    fp._p.append(fld)
    sectPr = sec._sectPr
    pg = sectPr.find(qn("w:pgNumType"))
    if pg is None:
        pg = OxmlElement("w:pgNumType"); sectPr.append(pg)
    if roman:
        pg.set(qn("w:fmt"), "lowerRoman")
    if start is not None:
        pg.set(qn("w:start"), str(start))

if ch1p is not None and len(doc.sections) == 1:
    # 在第一章標題段前一段尾插入 sectPr（結束前文節）
    prev = ch1p._p.getprevious()
    if prev is not None and prev.tag == qn("w:p"):
        ppr = prev.find(qn("w:pPr"))
        if ppr is None:
            ppr = OxmlElement("w:pPr"); prev.insert(0, ppr)
        if ppr.find(qn("w:sectPr")) is None:
            body_sectPr = doc.sections[-1]._sectPr
            front = copy.deepcopy(body_sectPr)
            # 移除 front 既有 headerReference/footerReference 以便獨立
            ppr.append(front)
    print("S2 section break inserted:", len(doc.sections), "sections")

if len(doc.sections) >= 2:
    set_footer_pagenum(doc.sections[0], roman=True, start=1)
    set_footer_pagenum(doc.sections[1], start=1)
    doc.sections[0].different_first_page_header_footer = True   # 封面不顯示頁碼
    print("S2 page numbers: front=lowerRoman (cover blank), body=arabic restart")

doc.save(str(TGT))
mp = sum(1 for p in doc.paragraphs if len(list(p._p.iter(M))))
print(f"saved. math-paras={mp} sections={len(doc.sections)} paras={len(doc.paragraphs)} "
      f"images={len(doc.inline_shapes)} tables={len(doc.tables)}")

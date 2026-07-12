#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
assemble_thesis.py -- merge draft/ch1..ch6.md into 260710 碩論_全文改寫_claude.docx
Conventions handled:
  * '# / ## / ###' -> Heading 1/2/3
  * body paragraph ending with 【tag】 -> tag stripped and re-appended as (tag) with
    YELLOW highlight (Wade's paragraph-function habit)
  * '> 插圖：圖5-XX ...' -> embed result/圖5-XX_*.png (15.5 cm) + centred caption
  * '> 插表：表5-X ...' -> embed result/表5-X_*.png + caption ABOVE
  * '> 原圖保留：...' -> gray italic placeholder line (original figure kept in Wade's doc)
Abstract ('# 摘要') from ch6.md is placed first.
"""
import glob
import re
from pathlib import Path
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Pt, Cm, RGBColor

HERE = Path(__file__).parent
DRAFT = HERE / "draft"
RESULT = HERE / "result"
OUT = HERE / "260711_TX碩論_Wade.docx"

doc = docx.Document()
st = doc.styles["Normal"]
st.font.name = "Times New Roman"
st.element.rPr.rFonts.set(docx.oxml.ns.qn("w:eastAsia"), "標楷體")
st.font.size = Pt(12)
doc.styles["Heading 1"].font.size = Pt(18)
doc.styles["Heading 2"].font.size = Pt(15)
doc.styles["Heading 3"].font.size = Pt(13)

TAG_RE = re.compile(r"【([^】]{2,10})】\s*$")

def add_body(text):
    m = TAG_RE.search(text)
    tag = None
    if m:
        tag = m.group(1)
        text = text[: m.start()].rstrip()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.85)
    p.paragraph_format.line_spacing = 1.5
    p.add_run(text)
    if tag:
        r = p.add_run(f"({tag})")
        r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return p

def add_caption(text, above=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.bold = True
    return p

def find_asset(prefix):
    hits = sorted(glob.glob(str(RESULT / f"{prefix}*.png")))
    return hits[0] if hits else None

def add_figure(line):
    m = re.search(r"圖\s*(\d+)[-.]0?(\d+)", line)
    cap = line.split("：", 1)[-1].strip()
    if m:
        ch, num = m.group(1), int(m.group(2))
        cands = []
        if ch == "5":
            a = find_asset(f"圖5-{num:02d}")
            if a: cands = [a]
        if not cands:
            cands = (sorted(glob.glob(str(HERE / "thesis_src" / "origfigs" / f"圖{ch}-{num}.*"))) or
                     sorted(glob.glob(str(HERE / "thesis_src" / "origfigs" / f"圖{ch}-{num}_*.*"))))
        ok = False
        for img in cands:
            try:
                doc.add_picture(img, width=Cm(15.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                ok = True
            except Exception:
                pass
        if ok:
            add_caption(cap)
            return
    add_caption(f"〔插圖佔位〕{cap}")

def add_table_png(line):
    m = re.search(r"表5-(\d+)", line)
    cap = line.split("：", 1)[-1].strip()
    add_caption(cap, above=True)
    if m:
        img = find_asset(f"表5-{m.group(1)}")
        if img:
            doc.add_picture(img, width=Cm(15.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            return
    add_caption(f"〔插表佔位〕{cap}")

ORIG = HERE / "thesis_src" / "origfigs"

def add_placeholder(line):
    # 原圖保留：圖X-X 說明  -> embed extracted original image(s) + caption
    m = re.search(r"圖\s*(\d+)[-.](\d+)", line)
    cap = line.split("：", 1)[-1].strip()
    if m and line.strip().startswith("原圖"):
        key = f"圖{m.group(1)}-{int(m.group(2))}"
        hits = sorted(glob.glob(str(ORIG / f"{key}.*"))) or                sorted(glob.glob(str(ORIG / f"{key}_*.*")))
        if hits:
            ok = False
            for img in hits:
                try:
                    doc.add_picture(img, width=Cm(14.5))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    ok = True
                except Exception:
                    pass
            if ok:
                add_caption(cap)
                return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"〔{line.strip()}〕")
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    r.italic = True

def process(md_text):
    for raw in md_text.split("\n"):
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith(">"):
            body = line.lstrip("> ").strip()
            if body.startswith("插圖"):
                add_figure(body)
            elif body.startswith("插表"):
                add_table_png(body)
            else:
                add_placeholder(body)
        else:
            add_body(line.strip())

# abstract first (ch6.md contains 摘要 then 第六章)
ch6 = (DRAFT / "ch6.md").read_text(encoding="utf-8")
if "# 第六章" in ch6:
    abstract, ch6_body = ch6.split("# 第六章", 1)
    ch6_body = "# 第六章" + ch6_body
else:
    abstract, ch6_body = "", ch6
process(abstract)
doc.add_page_break()
for n in (1, 2, 3, 4, 5):
    process((DRAFT / f"ch{n}.md").read_text(encoding="utf-8"))
    doc.add_page_break()
process(ch6_body)

# ---- font enforcement: Times New Roman + 標楷體 on every run ----
def enforce(runs):
    for r in runs:
        r.font.name = "Times New Roman"
        rpr = r._r.get_or_add_rPr()
        rf = rpr.find(docx.oxml.ns.qn("w:rFonts"))
        if rf is None:
            rf = docx.oxml.OxmlElement("w:rFonts")
            rpr.append(rf)
        rf.set(docx.oxml.ns.qn("w:eastAsia"), "標楷體")
for p in doc.paragraphs:
    enforce(p.runs)
for t in doc.tables:
    for row in t.rows:
        for c in row.cells:
            for p in c.paragraphs:
                enforce(p.runs)

doc.save(str(OUT))
print("saved", OUT.name, "paragraphs:", len(doc.paragraphs))
